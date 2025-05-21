# Enhanced create_chroma_db.py for medical data dictionary
# Comprehensive processing for numbered definitions, embedded XLS, and table content

import os
import shutil
import re
from pathlib import Path
from langchain_community.document_loaders import (
    UnstructuredExcelLoader, 
    TextLoader,
    UnstructuredWordDocumentLoader,
    PyPDFLoader
)
from langchain_unstructured import UnstructuredLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_vertexai import VertexAIEmbeddings
from langchain.schema import Document
import chromadb
import logging
import json
from typing import List, Dict, Any, Optional, Tuple, Union
import pandas as pd
import numpy as np
from docx import Document as DocxDocument
from docx.table import Table as DocxTable

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
DOCUMENTS_DIR = "./documents" 
CHROMA_DB_DIR = "./chroma_db_data"
COLLECTION_NAME = "rag_collection"
EMBEDDING_MODEL_NAME = "text-embedding-005"

# Enhanced chunking for complete numbered definitions and tables
CHUNK_SIZE = 3000  # Larger for complete definitions
CHUNK_OVERLAP = 500  # More overlap for context preservation
MIN_DEFINITION_SIZE = 50  # Minimum size for definition chunks
MAX_DEFINITION_SIZE = 2500  # Maximum size before forced splitting
TABLE_CHUNK_SIZE = 2000  # Special size for tables

# Initialize Vertex AI Embeddings
try:
    embeddings = VertexAIEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    logger.info(f"Successfully initialized VertexAIEmbeddings: {EMBEDDING_MODEL_NAME}")
except Exception as e:
    logger.error(f"Error initializing VertexAIEmbeddings: {e}")
    logger.info("Please ensure Google Cloud authentication and Vertex AI API access")
    exit(1)

def preprocess_text(text_content: str) -> str:
    """Enhanced text preprocessing for medical documentation"""
    if not text_content:
        return ""
    
    # Remove extra whitespace but preserve structure
    text_content = re.sub(r' +', ' ', text_content)
    text_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', text_content)
    
    # Clean non-printable characters but keep medical symbols and formatting
    text_content = re.sub(r'[^\x20-\x7E\n\r\t•\-–\—]', ' ', text_content)
    
    # Fix common OCR/extraction issues
    text_content = re.sub(r'(\d+)\s*\.\s*([A-Z])', r'\1. \2', text_content)  # Fix numbered definitions
    text_content = re.sub(r'([a-z])\s*:\s*([A-Z])', r'\1: \2', text_content)  # Fix field separators
    
    # Preserve bullet points and formatting
    text_content = re.sub(r'[•·]\s*', '• ', text_content)  # Normalize bullet points
    text_content = re.sub(r'\*\*([^*]+)\*\*', r'**\1**', text_content)  # Fix bold formatting
    
    return text_content.strip()

def detect_numbered_definition_boundaries(text: str) -> List[Tuple[int, int, str, str]]:
    """Detect numbered definition boundaries with improved pattern matching"""
    boundaries = []
    
    # Enhanced patterns for numbered definitions
    patterns = [
        r'(\d+)\.\s*([^:\n]+):\s*',  # Standard: "158. Diagnosis Code 6:"
        r'\*\*(\d+)\.\s*([^:*\n]+):\*\*',  # Bold: "**158. Diagnosis Code 6:**"
        r'(\d+)\.\s*\[([^\]]+)\]\s*:',  # Bracketed: "158. [Diagnosis Code 6]:"
        r'(\d+)\.\s*([A-Z][^:\n]*[a-z])\s*:',  # Title case: "158. Diagnosis Code 6:"
    ]
    
    all_matches = []
    
    for pattern in patterns:
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        for match in matches:
            start_pos = match.start()
            definition_number = match.group(1)
            definition_title = match.group(2).strip()
            all_matches.append((start_pos, definition_number, definition_title, match.group(0)))
    
    # Sort by position and remove duplicates
    all_matches.sort(key=lambda x: x[0])
    
    # Remove overlapping matches
    final_matches = []
    last_end = -1
    
    for start_pos, number, title, full_match in all_matches:
        if start_pos > last_end:
            final_matches.append((start_pos, number, title, full_match))
            last_end = start_pos + len(full_match)
    
    # Convert to boundaries with end positions
    for i, (start_pos, number, title, full_match) in enumerate(final_matches):
        # Find the end of this definition
        if i + 1 < len(final_matches):
            end_pos = final_matches[i + 1][0]
        else:
            # Look for next major section or end of text
            next_section = re.search(r'\n\n[A-Z]', text[start_pos:])
            if next_section:
                end_pos = start_pos + next_section.start()
            else:
                end_pos = len(text)
        
        boundaries.append((start_pos, end_pos, f"{number}. {title}", number))
    
    return boundaries

def extract_docx_tables(file_path: str) -> List[Dict]:
    """Extract tables directly from DOCX files including embedded XLS content"""
    tables = []
    logger.info(f"Extracting tables from DOCX file: {file_path}")
    
    try:
        doc = DocxDocument(file_path)
        
        for table_idx, table in enumerate(doc.tables):
            logger.info(f"Processing table {table_idx + 1} of {len(doc.tables)}")
            table_data = []
            headers = []
            
            # Extract table headers
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                logger.info(f"  Table {table_idx + 1}: Found {len(headers)} headers")
            
            # Extract table data
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_data = [cell.text.strip() for cell in row.cells]
                if any(cell.strip() for cell in row_data):  # Skip empty rows
                    table_data.append(row_data)
            
            if table_data or headers:
                # Format table as text
                formatted_content = []
                if headers:
                    formatted_content.append(" | ".join(headers))
                    formatted_content.append("-" * len(" | ".join(headers)))
                
                for row in table_data[:50]:  # Limit rows for performance
                    formatted_content.append(" | ".join(row))
                
                tables.append({
                    'type': 'docx_table',
                    'table_index': table_idx,
                    'headers': headers,
                    'data': table_data,
                    'formatted_content': '\n'.join(formatted_content),
                    'row_count': len(table_data),
                    'col_count': len(headers) if headers else 0
                })
                
                logger.info(f"  Successfully extracted DOCX table {table_idx + 1}: {len(table_data)} rows, {len(headers)} columns")
        
    except Exception as e:
        logger.error(f"Error extracting DOCX tables: {e}")
    
    logger.info(f"Total DOCX tables extracted: {len(tables)}")
    return tables

def extract_embedded_tables_advanced(content: str) -> List[Dict]:
    """Advanced table extraction with multiple detection methods"""
    tables = []
    
    # Method 1: Value mapping tables (enhanced patterns)
    value_patterns = [
        r'([A-Z0-9]+)\s*[=\-]\s*([^\n]+)',  # A = Description
        r'([A-Z0-9]+)\s+([A-Z][^\n]+)',     # A Description
        r'(\d+)\s+([A-Z][^\n]+)',           # 1 Description
    ]
    
    for pattern in value_patterns:
        value_matches = re.findall(pattern, content)
        if len(value_matches) > 3:  # Must have multiple entries
            tables.append({
                'type': 'value_mapping',
                'data': value_matches,
                'formatted_content': '\n'.join([f"{code} = {desc}" for code, desc in value_matches]),
                'entry_count': len(value_matches)
            })
            break
    
    # Method 2: Field specification tables
    field_specs = re.findall(r'(Format|Technical Name|Length|Position[s]?):\s*([^\n]+)', content)
    if field_specs:
        tables.append({
            'type': 'field_specification',
            'data': field_specs,
            'formatted_content': '\n'.join([f"• **{field}:** {value}" for field, value in field_specs]),
            'field_count': len(field_specs)
        })
    
    # Method 3: Appendix/lookup tables with headers
    lines = content.split('\n')
    table_blocks = []
    current_block = []
    in_table = False
    
    for line in lines:
        # Detect table start (headers with separators or consistent formatting)
        if (re.match(r'^[A-Z][^|]*\|[^|]*', line) or 
            re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+', line) or
            '---' in line or '===' in line):
            if not in_table:
                in_table = True
                current_block = [line]
            else:
                current_block.append(line)
        elif in_table:
            if line.strip() and not re.match(r'^\d+\.', line):  # Continue table
                current_block.append(line)
            else:  # End of table
                if len(current_block) > 2:
                    table_blocks.append(current_block)
                current_block = []
                in_table = False
    
    # Process detected table blocks
    for i, block in enumerate(table_blocks):
        cleaned_block = [line.strip() for line in block if line.strip()]
        if len(cleaned_block) > 2:
            # Try to detect headers
            headers = []
            data_rows = cleaned_block
            
            # Check if first row looks like headers
            first_row = cleaned_block[0]
            if ('|' in first_row or 
                re.match(r'^[A-Z][^a-z]*\s+[A-Z][^a-z]*', first_row)):
                headers = re.split(r'\s*\|\s*', first_row) if '|' in first_row else first_row.split()
                data_rows = cleaned_block[1:]
            
            tables.append({
                'type': 'structured_table',
                'table_index': i,
                'headers': headers,
                'data_rows': data_rows,
                'formatted_content': '\n'.join(cleaned_block),
                'row_count': len(data_rows)
            })
    
    # Method 4: Code tables from appendices
    appendix_pattern = r'([A-Z0-9]{1,4})\s+([\w\s\-\']+?)(?=\n[A-Z0-9]{1,4}\s+|\n\n|\Z)'
    appendix_matches = re.findall(appendix_pattern, content, re.MULTILINE)
    
    if len(appendix_matches) > 5:  # Substantial lookup table
        tables.append({
            'type': 'appendix_lookup',
            'data': appendix_matches,
            'formatted_content': '\n'.join([f"{code:5s} {desc}" for code, desc in appendix_matches]),
            'entry_count': len(appendix_matches)
        })
    
    return tables

def analyze_content_patterns_advanced(text: str) -> Dict[str, Any]:
    """Advanced content pattern analysis with enhanced detection"""
    patterns = {}
    
    # 1. Numbered definitions (enhanced detection)
    numbered_def_patterns = [
        r'(\d+)\.\s*([^:\n]+):',
        r'\*\*(\d+)\.\s*([^:*\n]+):\*\*',
        r'(\d+)\.\s*\[([^\]]+)\]:',
    ]
    
    all_numbered_defs = []
    for pattern in numbered_def_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        all_numbered_defs.extend(matches)
    
    # Remove duplicates and sort
    unique_defs = list(set(all_numbered_defs))
    patterns['numbered_definitions'] = unique_defs
    patterns['definition_count'] = len(unique_defs)
    
    # Extract definition numbers for reference
    definition_numbers = sorted(set([match[0] for match in unique_defs if match[0].isdigit()]))
    patterns['definition_numbers'] = definition_numbers
    
    # 2. Technical names (enhanced patterns)
    tech_name_patterns = [
        r'Technical Name:\s*([A-Z_][A-Z0-9_]*)',
        r'Technical Name:\s*([A-Z][A-Z0-9\s_]*)',
        r'\*\*Technical Name:\*\*\s*([A-Z_][A-Z0-9_]*)',
    ]
    
    technical_names = []
    for pattern in tech_name_patterns:
        matches = re.findall(pattern, text)
        technical_names.extend(matches)
    
    patterns['technical_names'] = list(set(technical_names))
    
    # 3. Field specifications (comprehensive)
    field_elements = {
        'formats': re.findall(r'[*•]\s*\*\*Format:\*\*\s*([^\n]+)', text),
        'lengths': re.findall(r'[*•]\s*\*\*Length:\*\*\s*(\d+(?:\.\d+)?)', text),
        'positions': re.findall(r'[*•]\s*\*\*Position[s]?:\*\*\s*([\d\s\-–—,]+)', text),
        'definitions': re.findall(r'[*•]\s*\*\*Definition:\*\*\s*([^\n]+)', text)
    }
    
    patterns.update(field_elements)
    
    # 4. Value mappings and codes
    value_mappings = []
    value_patterns = [
        r'([A-Z0-9]{1,4})\s*=\s*([^\n]+)',
        r'([A-Z0-9]{1,4})\s+([A-Z][^\n]+)',
        r'(\d+)\s+([A-Z][A-Za-z\s]+?)(?=\n\d+|\n\n|\Z)'
    ]
    
    for pattern in value_patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        if len(matches) > 3:  # Substantial list
            value_mappings.extend(matches)
    
    patterns['value_mappings'] = value_mappings[:50]  # Limit for performance
    
    # 5. Cross-references and relationships
    field_refs = re.findall(r'[Ff]ield\s*#?\s*(\d+)', text)
    appendix_refs = len(re.findall(r'[Aa]ppendix|[Aa]ppendices', text))
    see_also_refs = re.findall(r'[Ss]ee\s+(?:also\s+)?[Ff]ield\s*#?(\d+)', text)
    
    patterns['field_references'] = field_refs
    patterns['see_also_references'] = see_also_refs
    patterns['has_appendix_refs'] = appendix_refs > 0
    patterns['appendix_ref_count'] = appendix_refs
    
    # 6. Content classification (enhanced logic)
    content_type = 'descriptive_text'  # Default
    
    if unique_defs:
        if len(unique_defs) == 1:
            content_type = 'single_numbered_definition'
        else:
            content_type = 'multiple_numbered_definitions'
    elif technical_names and any(field_elements.values()):
        content_type = 'field_specification'
    elif len(value_mappings) > 10:
        content_type = 'value_mapping_table'
    elif appendix_refs > 0 and len(value_mappings) > 3:
        content_type = 'appendix_table'
    elif 'Aetna' in text and ('Field' in text or 'Definition' in text):
        content_type = 'medical_data_dictionary'
    
    patterns['content_type'] = content_type
    
    # 7. Quality indicators
    patterns['has_complete_structure'] = (
        len(field_elements['formats']) > 0 and
        len(patterns['technical_names']) > 0 and
        len(field_elements['definitions']) > 0
    )
    
    patterns['completeness_score'] = sum([
        1 if field_elements['formats'] else 0,
        1 if patterns['technical_names'] else 0,
        1 if field_elements['lengths'] else 0,
        1 if field_elements['positions'] else 0,
        1 if field_elements['definitions'] else 0
    ])
    
    # 8. Table indicators
    patterns['has_tabular_data'] = (
        '|' in text or 
        len(value_mappings) > 5 or
        re.search(r'\n[A-Z0-9]+\s+[A-Z]', text) is not None
    )
    
    return patterns

def create_enhanced_metadata(doc: Document, patterns: Dict, tables: List[Dict], file_path: str, chunk_info: Dict = None) -> Dict:
    """Create comprehensive metadata for ChromaDB with enhanced field detection"""
    metadata = {
        'source_filename': os.path.basename(file_path),
        'content_length': len(doc.page_content),
        'content_type': patterns.get('content_type', 'unknown'),
    }
    
    # Add original document metadata
    if hasattr(doc, 'metadata') and doc.metadata:
        metadata.update({
            'page_number': str(doc.metadata.get('page_number', 'N/A')),
            'element_type': doc.metadata.get('category', doc.metadata.get('element_type', 'Unknown')),
        })
    else:
        metadata.update({
            'page_number': 'N/A',
            'element_type': 'Unknown',
        })
    
    # Chunk information
    if chunk_info:
        metadata.update(chunk_info)
    
    # Numbered definition metadata (enhanced)
    numbered_defs = patterns.get('numbered_definitions', [])
    if numbered_defs:
        # Handle both single and multiple definitions
        definition_numbers = [def_info[0] for def_info in numbered_defs if len(def_info) > 0]
        definition_titles = [def_info[1] for def_info in numbered_defs if len(def_info) > 1]
        
        metadata['definition_numbers'] = ','.join(definition_numbers)
        metadata['definition_titles'] = ','.join(definition_titles)
        metadata['is_numbered_definition'] = True
        metadata['definition_count'] = len(numbered_defs)
        
        # Single definition metadata
        if len(numbered_defs) == 1:
            metadata['definition_number'] = definition_numbers[0]
            metadata['definition_title'] = definition_titles[0] if definition_titles else ''
        
        # Check for specific fields
        for number, title in numbered_defs:
            if 'diagnosis code' in title.lower():
                metadata['is_diagnosis_field'] = True
                metadata['diagnosis_type'] = title
            elif 'hcfa' in title.lower():
                metadata['is_hcfa_field'] = True
            elif 'admit' in title.lower():
                metadata['is_admit_field'] = True
    else:
        metadata['is_numbered_definition'] = False
        metadata['definition_count'] = 0
    
    # Technical name metadata (enhanced)
    technical_names = patterns.get('technical_names', [])
    if technical_names:
        metadata['technical_names'] = ','.join(technical_names)
        metadata['is_field_spec'] = True
        metadata['technical_name_count'] = len(technical_names)
        
        if len(technical_names) == 1:
            metadata['primary_technical_name'] = technical_names[0]
            
            # Classify technical name types
            tech_name = technical_names[0]
            if 'DX_CD' in tech_name:
                metadata['field_category'] = 'diagnosis'
            elif 'AMT' in tech_name:
                metadata['field_category'] = 'amount'
            elif 'CD' in tech_name:
                metadata['field_category'] = 'code'
            elif 'DT' in tech_name:
                metadata['field_category'] = 'date'
            elif 'ID' in tech_name:
                metadata['field_category'] = 'identifier'
    else:
        metadata['is_field_spec'] = False
        metadata['technical_name_count'] = 0
    
    # Field specification elements
    field_elements = ['formats', 'lengths', 'positions', 'definitions']
    for element in field_elements:
        values = patterns.get(element, [])
        if values:
            metadata[f'{element[:-1]}_info'] = ','.join(str(v) for v in values[:3])  # Limit for metadata size
    
    # Content quality and completeness
    completeness_score = patterns.get('completeness_score', 0)
    metadata['definition_completeness'] = completeness_score
    metadata['is_complete_definition'] = completeness_score >= 4
    metadata['has_complete_structure'] = patterns.get('has_complete_structure', False)
    
    # Table and structured data metadata
    metadata['has_embedded_tables'] = len(tables) > 0
    metadata['table_count'] = len(tables)
    metadata['has_tabular_data'] = patterns.get('has_tabular_data', False)
    
    if tables:
        table_types = list(set([table['type'] for table in tables]))
        metadata['table_types'] = ','.join(table_types)
        
        # Specific table information
        total_entries = sum(table.get('entry_count', table.get('row_count', 0)) for table in tables)
        metadata['total_table_entries'] = total_entries
    
    # Value mapping metadata
    value_mappings = patterns.get('value_mappings', [])
    metadata['has_value_mappings'] = len(value_mappings) > 0
    metadata['value_mapping_count'] = len(value_mappings)
    
    # Reference and relationship metadata
    field_refs = patterns.get('field_references', [])
    see_also_refs = patterns.get('see_also_references', [])
    
    metadata['field_references'] = ','.join(field_refs) if field_refs else ''
    metadata['see_also_references'] = ','.join(see_also_refs) if see_also_refs else ''
    metadata['has_field_references'] = len(field_refs) > 0
    metadata['has_see_also_references'] = len(see_also_refs) > 0
    metadata['has_appendix_references'] = patterns.get('has_appendix_refs', False)
    metadata['appendix_ref_count'] = patterns.get('appendix_ref_count', 0)
    
    # Classification indicators
    metadata['is_appendix'] = patterns.get('content_type') in ['appendix_table', 'value_mapping_table']
    metadata['is_lookup_table'] = patterns.get('content_type') == 'value_mapping_table'
    metadata['is_single_definition'] = patterns.get('content_type') == 'single_numbered_definition'
    metadata['is_multiple_definitions'] = patterns.get('content_type') == 'multiple_numbered_definitions'
    
    # Search optimization metadata
    searchable_terms = []
    
    # Add numbered definition terms
    if numbered_defs:
        for number, title in numbered_defs:
            searchable_terms.extend([
                f"{number}. {title}",
                title.upper(),
                title.lower(),
                title.replace(' ', '_').upper(),  # Technical name format
                number  # Just the number
            ])
    
    # Add technical names
    if technical_names:
        searchable_terms.extend(technical_names)
        searchable_terms.extend([name.lower() for name in technical_names])
    
    # Add field references
    if field_refs:
        searchable_terms.extend([f"Field {ref}" for ref in field_refs])
    
    # Limit searchable terms for ChromaDB constraints
    if searchable_terms:
        unique_terms = list(dict.fromkeys(searchable_terms))  # Remove duplicates while preserving order
        metadata['searchable_terms'] = ','.join(unique_terms[:20])
    
    # Content categorization for better retrieval
    content_categories = []
    if metadata.get('is_numbered_definition'):
        content_categories.append('numbered_definition')
    if metadata.get('is_field_spec'):
        content_categories.append('field_specification')
    if metadata.get('has_embedded_tables'):
        content_categories.append('contains_tables')
    if metadata.get('is_appendix'):
        content_categories.append('appendix')
    if metadata.get('has_value_mappings'):
        content_categories.append('value_mappings')
    
    metadata['content_categories'] = ','.join(content_categories) if content_categories else ''
    
    # Ensure all metadata values are ChromaDB compatible
    for key, value in metadata.items():
        if isinstance(value, (list, tuple)):
            metadata[key] = ','.join(str(v) for v in value)
        elif not isinstance(value, (str, int, float, bool, type(None))):
            metadata[key] = str(value)
        elif value is None:
            metadata[key] = ''
        # Handle very long strings
        elif isinstance(value, str) and len(value) > 500:
            metadata[key] = value[:500] + '...'
    
    return metadata

def smart_chunk_numbered_definitions(text: str, text_splitter: RecursiveCharacterTextSplitter) -> List[Document]:
    """Intelligently chunk text while preserving complete numbered definitions"""
    boundaries = detect_numbered_definition_boundaries(text)
    chunks = []
    
    if not boundaries:
        # No numbered definitions found, use standard chunking
        return text_splitter.create_documents([text])
    
    current_pos = 0
    
    for start_pos, end_pos, definition_title, definition_number in boundaries:
        # Add any text before the definition
        if current_pos < start_pos:
            pre_text = text[current_pos:start_pos].strip()
            if pre_text and len(pre_text) > MIN_DEFINITION_SIZE:
                pre_chunks = text_splitter.create_documents([pre_text])
                chunks.extend(pre_chunks)
        
        # Extract the complete definition
        definition_text = text[start_pos:end_pos].strip()
        
        if len(definition_text) <= MAX_DEFINITION_SIZE:
            # Keep complete definition as one chunk
            chunk_doc = Document(
                page_content=definition_text,
                metadata={
                    'chunk_type': 'complete_definition',
                    'definition_number': definition_number,
                    'definition_title': definition_title.replace(f"{definition_number}. ", ""),
                    'start_pos': start_pos,
                    'end_pos': end_pos
                }
            )
            chunks.append(chunk_doc)
        else:
            # Definition is too long, split carefully at logical boundaries
            lines = definition_text.split('\n')
            current_chunk_lines = []
            current_length = 0
            header_lines = []
            
            # Extract header (first few lines with definition title)
            for i, line in enumerate(lines[:5]):
                if definition_number in line or definition_title.split()[-1] in line:
                    header_lines.append(line)
                else:
                    break
            
            chunk_index = 0
            for line in lines:
                # Check if adding this line exceeds size limit
                if current_length + len(line) > CHUNK_SIZE and current_chunk_lines:
                    # Create chunk with header if it's not the first chunk
                    chunk_content = '\n'.join(current_chunk_lines)
                    if chunk_index > 0 and header_lines:
                        chunk_content = '\n'.join(header_lines) + '\n...\n' + chunk_content
                    
                    chunk_doc = Document(
                        page_content=chunk_content,
                        metadata={
                            'chunk_type': 'partial_definition',
                            'definition_number': definition_number,
                            'definition_title': definition_title.replace(f"{definition_number}. ", ""),
                            'chunk_index': chunk_index,
                            'is_continuation': chunk_index > 0
                        }
                    )
                    chunks.append(chunk_doc)
                    
                    current_chunk_lines = []
                    current_length = 0
                    chunk_index += 1
                
                current_chunk_lines.append(line)
                current_length += len(line) + 1
            
            # Add final chunk
            if current_chunk_lines:
                chunk_content = '\n'.join(current_chunk_lines)
                if chunk_index > 0 and header_lines:
                    chunk_content = '\n'.join(header_lines) + '\n...\n' + chunk_content
                
                chunk_doc = Document(
                    page_content=chunk_content,
                    metadata={
                        'chunk_type': 'partial_definition',
                        'definition_number': definition_number,
                        'definition_title': definition_title.replace(f"{definition_number}. ", ""),
                        'chunk_index': chunk_index,
                        'is_continuation': chunk_index > 0,
                        'is_final_chunk': True
                    }
                )
                chunks.append(chunk_doc)
        
        current_pos = end_pos
    
    # Add any remaining text
    if current_pos < len(text):
        remaining_text = text[current_pos:].strip()
        if remaining_text and len(remaining_text) > MIN_DEFINITION_SIZE:
            remaining_chunks = text_splitter.create_documents([remaining_text])
            chunks.extend(remaining_chunks)
    
    return chunks

def load_document_with_comprehensive_extraction(file_path: str) -> List[Document]:
    """Load document with comprehensive extraction of content, tables, and metadata"""
    extension = os.path.splitext(file_path)[1].lower()
    documents = []
    
    logger.info(f"Loading document: {file_path} ({extension})")
    
    # Choose appropriate loader and extract method
    try:
        if extension == ".docx":
            # Use both unstructured loader and direct DOCX processing
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            raw_docs = loader.load()
            
            # Also extract tables directly from DOCX
            docx_tables = extract_docx_tables(file_path)
            
        elif extension == ".doc":
            loader = UnstructuredLoader(file_path, mode="elements", strategy="auto")
            raw_docs = loader.load()
            docx_tables = []
            
        elif extension in [".xlsx", ".xls"]:
            # Handle Excel files specially
            return load_excel_comprehensive(file_path)
            
        elif extension == ".pdf":
            loader = PyPDFLoader(file_path)
            raw_docs = loader.load()
            docx_tables = []
            
        elif extension == ".txt":
            loader = TextLoader(file_path, encoding="utf-8")
            raw_docs = loader.load()
            docx_tables = []
            
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return documents
        
        # Process main content
        combined_content = ""
        for doc in raw_docs:
            doc.page_content = preprocess_text(doc.page_content)
            if len(doc.page_content.strip()) >= MIN_DEFINITION_SIZE:
                combined_content += doc.page_content + "\n\n"
        
        # Analyze patterns in combined content
        patterns = analyze_content_patterns_advanced(combined_content)
        embedded_tables = extract_embedded_tables_advanced(combined_content)
        
        # Combine extracted tables
        all_tables = embedded_tables + docx_tables
        
        # Create document with combined content
        if combined_content.strip():
            main_doc = Document(
                page_content=combined_content.strip(),
                metadata={
                    'source_filename': os.path.basename(file_path),
                    'element_type': 'combined_content',
                    'page_number': 'combined',
                }
            )
            
            # Create enhanced metadata
            enhanced_metadata = create_enhanced_metadata(main_doc, patterns, all_tables, file_path)
            main_doc.metadata = enhanced_metadata
            
            documents.append(main_doc)
        
        # Create separate documents for each extracted table
        for i, table in enumerate(all_tables):
            table_doc = Document(
                page_content=table['formatted_content'],
                metadata={
                    'source_filename': os.path.basename(file_path),
                    'element_type': f"table_{table['type']}",
                    'table_index': i,
                    'table_type': table['type'],
                    'is_embedded_table': True,
                    'is_table_content': True,
                    'table_entry_count': table.get('entry_count', table.get('row_count', 0)),
                    'page_number': 'table',
                }
            )
            
            # Add table-specific metadata
            table_metadata = create_enhanced_metadata(table_doc, {
                'content_type': 'table_content',
                'has_tabular_data': True,
                'value_mappings': table.get('data', [])
            }, [], file_path)
            table_doc.metadata.update(table_metadata)
            
            documents.append(table_doc)
    
    except Exception as e:
        logger.error(f"Error loading {file_path}: {e}")
        import traceback
        traceback.print_exc()
        
        # Try fallback loading for critical files
        try:
            if extension in [".docx", ".doc"]:
                loader = TextLoader(file_path, encoding="utf-8")
                fallback_docs = loader.load()
                for doc in fallback_docs:
                    doc.page_content = preprocess_text(doc.page_content)
                    doc.metadata['source_filename'] = os.path.basename(file_path)
                    doc.metadata['element_type'] = 'fallback_text'
                    documents.append(doc)
                logger.info(f"Used fallback loading for {file_path}")
        except Exception as fallback_error:
            logger.error(f"Fallback loading also failed: {fallback_error}")
    
    logger.info(f"Processed {len(documents)} elements from {file_path}")
    return documents

def load_excel_comprehensive(file_path: str) -> List[Document]:
    """Comprehensive Excel file processing with sheet and table extraction"""
    documents = []
    logger.info(f"Starting comprehensive Excel processing for: {file_path}")
    
    total_sheets_processed = 0
    successful_sheets = 0
    failed_sheets = 0
    
    try:
        # Method 1: Use pandas for structured data
        excel_file = pd.ExcelFile(file_path)
        total_sheets = len(excel_file.sheet_names)
        logger.info(f"Found {total_sheets} sheets in Excel file")
        
        for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
            logger.info(f"Processing sheet {sheet_idx + 1}/{total_sheets}: '{sheet_name}'")
            total_sheets_processed += 1
            
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                if df.empty:
                    logger.warning(f"  Sheet '{sheet_name}' is empty, skipping")
                    continue
                
                # Convert DataFrame to structured text
                content_lines = [f"**Sheet: {sheet_name}**", ""]
                
                # Add column headers
                headers = df.columns.tolist()
                content_lines.append("**Columns:**")
                content_lines.append(" | ".join(str(col) for col in headers))
                content_lines.append("-" * len(" | ".join(str(col) for col in headers)))
                content_lines.append("")
                
                # Add data rows (limit for performance)
                content_lines.append("**Data:**")
                for idx, row in df.head(100).iterrows():
                    row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
                    content_lines.append(row_text)
                
                # Analyze the sheet content
                sheet_content = "\n".join(content_lines)
                
                # Check if this looks like a lookup table
                is_lookup = False
                has_codes = False
                
                if len(df) > 5 and len(df.columns) >= 2:
                    # Check first column for codes/values
                    first_col = df.iloc[:, 0].astype(str)
                    if first_col.str.match(r'^[A-Z0-9]+$').any():
                        has_codes = True
                        is_lookup = True
                
                # Create document with enhanced metadata
                sheet_doc = Document(
                    page_content=sheet_content,
                    metadata={
                        'source_filename': os.path.basename(file_path),
                        'sheet_name': sheet_name,
                        'element_type': 'excel_sheet',
                        'is_embedded_table': True,
                        'is_table_content': True,
                        'table_type': 'excel_data',
                        'row_count': len(df),
                        'column_count': len(df.columns),
                        'has_headers': True,
                        'is_lookup_table': is_lookup,
                        'has_codes': has_codes,
                        'content_type': 'value_mapping_table' if is_lookup else 'data_table'
                    }
                )
                
                # Add patterns analysis
                patterns = {
                    'content_type': 'value_mapping_table' if is_lookup else 'data_table',
                    'has_tabular_data': True,
                    'value_mappings': []
                }
                
                if is_lookup and len(df.columns) >= 2:
                    # Extract value mappings
                    for _, row in df.head(50).iterrows():
                        code = str(row.iloc[0]) if pd.notna(row.iloc[0]) else ""
                        desc = str(row.iloc[1]) if pd.notna(row.iloc[1]) else ""
                        if code and desc:
                            patterns['value_mappings'].append((code, desc))
                
                enhanced_metadata = create_enhanced_metadata(sheet_doc, patterns, [], file_path)
                sheet_doc.metadata.update(enhanced_metadata)
                
                documents.append(sheet_doc)
                successful_sheets += 1
                
                logger.info(f"  Successfully processed sheet '{sheet_name}': {len(df)} rows, {len(df.columns)} columns")
                if is_lookup:
                    logger.info(f"    Detected as lookup table with {len(patterns.get('value_mappings', []))} value mappings")
                
            except Exception as e:
                failed_sheets += 1
                logger.warning(f"  Could not process sheet '{sheet_name}': {e}")
    
    except Exception as e:
        logger.error(f"Error processing Excel file {file_path}: {e}")
        
        # Fallback to unstructured loader
        try:
            logger.info("Attempting fallback Excel loading with UnstructuredExcelLoader...")
            loader = UnstructuredExcelLoader(file_path, mode="elements")
            fallback_docs = loader.load()
            for doc in fallback_docs:
                doc.page_content = preprocess_text(doc.page_content)
                doc.metadata['source_filename'] = os.path.basename(file_path)
                doc.metadata['element_type'] = 'excel_fallback'
                doc.metadata['is_table_content'] = True
                documents.append(doc)
            logger.info(f"Fallback Excel loading successful: {len(fallback_docs)} documents")
        except Exception as fallback_error:
            logger.error(f"Fallback Excel loading also failed: {fallback_error}")
    
    # Final logging summary
    logger.info(f"Excel processing complete:")
    logger.info(f"  Total sheets processed: {total_sheets_processed}")
    logger.info(f"  Successful sheets: {successful_sheets}")
    logger.info(f"  Failed sheets: {failed_sheets}")
    logger.info(f"  Total documents created: {len(documents)}")
    
    return documents

def load_all_documents(source_dir: str) -> List[Document]:
    """Load all documents from source directory with comprehensive processing"""
    all_documents = []
    excel_files_processed = 0
    total_excel_sheets = 0
    
    if not os.path.exists(source_dir):
        logger.error(f"Source directory '{source_dir}' not found")
        return all_documents
    
    supported_extensions = {'.docx', '.doc', '.xlsx', '.xls', '.pdf', '.txt'}
    
    for filename in os.listdir(source_dir):
        if filename.startswith('.'):
            continue
        
        file_path = os.path.join(source_dir, filename)
        extension = os.path.splitext(filename)[1].lower()
        
        if extension in supported_extensions and os.path.isfile(file_path):
            docs = load_document_with_comprehensive_extraction(file_path)
            all_documents.extend(docs)
            
            # Track Excel file processing
            if extension in ['.xlsx', '.xls']:
                excel_files_processed += 1
                # Count sheets processed
                excel_sheets = sum(1 for doc in docs if doc.metadata.get('element_type') == 'excel_sheet')
                total_excel_sheets += excel_sheets
                logger.info(f"  Excel summary: {excel_sheets} sheets processed from {filename}")
            
            # Log extraction results
            numbered_defs = sum(1 for doc in docs if doc.metadata.get('is_numbered_definition'))
            tables = sum(1 for doc in docs if doc.metadata.get('is_embedded_table'))
            complete_defs = sum(1 for doc in docs if doc.metadata.get('is_complete_definition'))
            
            logger.info(f"  Extracted: {numbered_defs} numbered definitions, {tables} tables/sheets, {complete_defs} complete definitions")
    
    # Final Excel processing summary
    if excel_files_processed > 0:
        logger.info(f"\n=== EXCEL PROCESSING SUMMARY ===")
        logger.info(f"Excel files processed: {excel_files_processed}")
        logger.info(f"Total Excel sheets processed: {total_excel_sheets}")
        logger.info(f"Average sheets per file: {total_excel_sheets/excel_files_processed:.1f}")
    
    logger.info(f"Total documents loaded: {len(all_documents)}")
    return all_documents

def create_optimized_chunks(documents: List[Document]) -> List[Document]:
    """Create optimized chunks with special handling for different content types"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=[
            "\n\n**",    # Definition headers
            "\n\n\n",    # Major section breaks
            "\n\n",      # Paragraph breaks
            "\n•",       # Bullet points
            "\n",        # Line breaks
            ". ",        # Sentence endings
            ": ",        # Definition separators
            "; ",        # Clause separators
            " ",         # Word boundaries
            ""           # Character level
        ]
    )
    
    final_chunks = []
    
    for doc in documents:
        content_type = doc.metadata.get('content_type', '')
        element_type = doc.metadata.get('element_type', '')
        
        # Handle different content types appropriately
        if doc.metadata.get('is_embedded_table') or doc.metadata.get('is_table_content'):
            # Keep tables intact
            doc.metadata['chunk_strategy'] = 'preserved_table'
            final_chunks.append(doc)
            
        elif content_type == 'single_numbered_definition' and len(doc.page_content) <= MAX_DEFINITION_SIZE:
            # Keep single complete definitions intact
            doc.metadata['chunk_strategy'] = 'preserved_complete_definition'
            final_chunks.append(doc)
            
        elif content_type in ['single_numbered_definition', 'multiple_numbered_definitions'] or doc.metadata.get('is_numbered_definition'):
            # Use smart chunking for numbered definitions
            definition_chunks = smart_chunk_numbered_definitions(doc.page_content, text_splitter)
            
            for i, chunk in enumerate(definition_chunks):
                # Preserve and enhance metadata
                chunk.metadata.update(doc.metadata)
                chunk.metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(definition_chunks),
                    'chunk_strategy': 'smart_numbered_definition'
                })
                final_chunks.append(chunk)
                
        elif content_type == 'value_mapping_table' and len(doc.page_content) <= TABLE_CHUNK_SIZE:
            # Keep value mapping tables intact
            doc.metadata['chunk_strategy'] = 'preserved_value_mapping'
            final_chunks.append(doc)
            
        else:
            # Standard chunking for other content
            chunks = text_splitter.split_documents([doc])
            
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_strategy': 'standard_recursive'
                })
                final_chunks.append(chunk)
    
    return final_chunks

def analyze_final_collection(documents: List[Document]) -> Dict[str, Any]:
    """Comprehensive analysis of the final document collection"""
    analysis = {
        'total_documents': len(documents),
        'numbered_definitions': 0,
        'single_definitions': 0,
        'multiple_definitions': 0,
        'complete_definitions': 0,
        'partial_definitions': 0,
        'embedded_tables': 0,
        'excel_sheets': 0,
        'value_mapping_tables': 0,
        'field_specifications': 0,
        'content_types': {},
        'chunk_strategies': {},
        'definition_numbers': set(),
        'technical_names': set(),
        'table_types': {},
        'quality_metrics': {},
        'field_categories': {},
        'excel_processing_stats': {
            'total_excel_docs': 0,
            'sheets_processed': 0,
            'lookup_tables': 0
        }
    }
    
    total_content_length = 0
    complexity_scores = []
    completeness_scores = []
    
    for doc in documents:
        metadata = doc.metadata
        
        # Count different types
        if metadata.get('is_numbered_definition'):
            analysis['numbered_definitions'] += 1
            
            # Track definition numbers
            if metadata.get('definition_number'):
                analysis['definition_numbers'].add(metadata['definition_number'])
            
            # Categorize by completeness
            if metadata.get('is_complete_definition'):
                analysis['complete_definitions'] += 1
            else:
                analysis['partial_definitions'] += 1
        
        # Track content type categories
        content_type = metadata.get('content_type', 'unknown')
        analysis['content_types'][content_type] = analysis['content_types'].get(content_type, 0) + 1
        
        if content_type == 'single_numbered_definition':
            analysis['single_definitions'] += 1
        elif content_type == 'multiple_numbered_definitions':
            analysis['multiple_definitions'] += 1
        
        # Track tables and Excel data
        if metadata.get('is_embedded_table'):
            analysis['embedded_tables'] += 1
            
            table_type = metadata.get('table_type', 'unknown')
            analysis['table_types'][table_type] = analysis['table_types'].get(table_type, 0) + 1
            
            if content_type == 'value_mapping_table':
                analysis['value_mapping_tables'] += 1
            
            # Track Excel-specific stats
            if metadata.get('element_type') == 'excel_sheet':
                analysis['excel_sheets'] += 1
                analysis['excel_processing_stats']['sheets_processed'] += 1
                if metadata.get('is_lookup_table'):
                    analysis['excel_processing_stats']['lookup_tables'] += 1
        
        # Track Excel documents
        if 'excel' in metadata.get('element_type', '').lower():
            analysis['excel_processing_stats']['total_excel_docs'] += 1
        
        # Track field specifications
        if metadata.get('is_field_spec'):
            analysis['field_specifications'] += 1
            
            # Track technical names
            if metadata.get('primary_technical_name'):
                analysis['technical_names'].add(metadata['primary_technical_name'])
            
            # Track field categories
            field_category = metadata.get('field_category', 'unknown')
            analysis['field_categories'][field_category] = analysis['field_categories'].get(field_category, 0) + 1
        
        # Track chunk strategies
        chunk_strategy = metadata.get('chunk_strategy', 'unknown')
        analysis['chunk_strategies'][chunk_strategy] = analysis['chunk_strategies'].get(chunk_strategy, 0) + 1
        
        # Collect metrics
        total_content_length += len(doc.page_content)
        
        complexity = metadata.get('complexity_score', 0)
        if isinstance(complexity, (int, float)):
            complexity_scores.append(complexity)
        
        completeness = metadata.get('definition_completeness', 0)
        if isinstance(completeness, (int, float)):
            completeness_scores.append(completeness)
    
    # Calculate quality metrics
    analysis['quality_metrics'] = {
        'average_content_length': total_content_length / len(documents) if documents else 0,
        'average_complexity': sum(complexity_scores) / len(complexity_scores) if complexity_scores else 0,
        'average_completeness': sum(completeness_scores) / len(completeness_scores) if completeness_scores else 0,
        'definition_coverage': len(analysis['definition_numbers']),
        'complete_definition_ratio': analysis['complete_definitions'] / max(analysis['numbered_definitions'], 1),
        'table_coverage': analysis['embedded_tables'],
        'excel_coverage': analysis['excel_sheets'],
        'technical_name_coverage': len(analysis['technical_names']),
        'content_diversity': len(analysis['content_types'])
    }
    
    # Convert sets to lists for JSON serialization
    analysis['definition_numbers'] = sorted(list(analysis['definition_numbers']))
    analysis['technical_names'] = sorted(list(analysis['technical_names']))
    
    return analysis

def main():
    """Main function to create comprehensive ChromaDB for medical data dictionary"""
    logger.info("Starting comprehensive ChromaDB creation...")
    logger.info("Enhanced for numbered definitions, embedded tables, and XLS processing")
    
    # Clean existing database
    if os.path.exists(CHROMA_DB_DIR):
        logger.info(f"Removing existing ChromaDB at {CHROMA_DB_DIR}")
        shutil.rmtree(CHROMA_DB_DIR)
    
    os.makedirs(CHROMA_DB_DIR, exist_ok=True)
    
    # Load and process documents
    logger.info("Loading documents with comprehensive extraction...")
    documents = load_all_documents(DOCUMENTS_DIR)
    
    if not documents:
        logger.error("No documents loaded. Check source directory and file formats.")
        return
    
    # Create optimized chunks
    logger.info("Creating optimized chunks...")
    chunked_docs = create_optimized_chunks(documents)
    
    # Analyze the collection
    analysis = analyze_final_collection(chunked_docs)
    
    # Log comprehensive analysis
    logger.info("\n=== COMPREHENSIVE COLLECTION ANALYSIS ===")
    logger.info(f"Total chunks: {analysis['total_documents']}")
    logger.info(f"Numbered definitions: {analysis['numbered_definitions']} ({analysis['single_definitions']} single, {analysis['multiple_definitions']} multiple)")
    logger.info(f"Complete definitions: {analysis['complete_definitions']} ({analysis['quality_metrics']['complete_definition_ratio']:.1%})")
    logger.info(f"Embedded tables: {analysis['embedded_tables']} (including {analysis['value_mapping_tables']} value mapping tables)")
    logger.info(f"Excel sheets processed: {analysis['excel_sheets']}")
    logger.info(f"Field specifications: {analysis['field_specifications']}")
    logger.info(f"Technical names: {analysis['quality_metrics']['technical_name_coverage']}")
    logger.info(f"Definition number coverage: {analysis['quality_metrics']['definition_coverage']} unique numbers")
    
    logger.info("\nContent Types Distribution:")
    for content_type, count in sorted(analysis['content_types'].items()):
        percentage = (count / analysis['total_documents']) * 100
        logger.info(f"  {content_type}: {count} ({percentage:.1f}%)")
    
    logger.info("\nChunk Strategies Used:")
    for strategy, count in sorted(analysis['chunk_strategies'].items()):
        percentage = (count / analysis['total_documents']) * 100
        logger.info(f"  {strategy}: {count} ({percentage:.1f}%)")
    
    if analysis['table_types']:
        logger.info("\nTable Types Found:")
        for table_type, count in sorted(analysis['table_types'].items()):
            logger.info(f"  {table_type}: {count}")
    
    if analysis['field_categories']:
        logger.info("\nField Categories:")
        for category, count in sorted(analysis['field_categories'].items()):
            logger.info(f"  {category}: {count}")
    
    # Excel processing summary
    excel_stats = analysis['excel_processing_stats']
    if excel_stats['total_excel_docs'] > 0:
        logger.info(f"\nExcel Processing Summary:")
        logger.info(f"  Excel documents: {excel_stats['total_excel_docs']}")
        logger.info(f"  Sheets processed: {excel_stats['sheets_processed']}")
        logger.info(f"  Lookup tables found: {excel_stats['lookup_tables']}")
    
    logger.info(f"\nQuality Metrics:")
    metrics = analysis['quality_metrics']
    logger.info(f"  Average content length: {metrics['average_content_length']:.0f} chars")
    logger.info(f"  Average completeness score: {metrics['average_completeness']:.1f}/5")
    logger.info(f"  Content diversity: {metrics['content_diversity']} types")
    
    # Prepare data for ChromaDB
    ids = [f"doc_{i}" for i in range(len(chunked_docs))]
    contents = [doc.page_content for doc in chunked_docs]
    metadatas = [doc.metadata for doc in chunked_docs]
    
    # Generate embeddings with progress tracking
    logger.info("Generating embeddings...")
    try:
        # Process in smaller batches for large datasets
        batch_size = 50
        embeddings_list = []
        
        for i in range(0, len(contents), batch_size):
            batch_contents = contents[i:i+batch_size]
            batch_embeddings = embeddings.embed_documents(batch_contents)
            embeddings_list.extend(batch_embeddings)
            
            progress = (i + len(batch_contents)) / len(contents) * 100
            logger.info(f"  Embedding progress: {progress:.1f}% ({i + len(batch_contents)}/{len(contents)})")
        
        logger.info(f"Generated {len(embeddings_list)} embeddings")
        
    except Exception as e:
        logger.error(f"Error generating embeddings: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # Initialize ChromaDB and store documents
    logger.info("Storing in ChromaDB...")
    try:
        client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
        collection = client.get_or_create_collection(name=COLLECTION_NAME)
        
        # Store in batches to handle large datasets
        batch_size = 500
        for i in range(0, len(ids), batch_size):
            end_idx = min(i + batch_size, len(ids))
            batch_ids = ids[i:end_idx]
            batch_embeddings = embeddings_list[i:end_idx]
            batch_contents = contents[i:end_idx]
            batch_metadatas = metadatas[i:end_idx]
            
            collection.add(
                ids=batch_ids,
                embeddings=batch_embeddings,
                documents=batch_contents,
                metadatas=batch_metadatas
            )
            
            progress = end_idx / len(ids) * 100
            logger.info(f"  Storage progress: {progress:.1f}% ({end_idx}/{len(ids)})")
        
        # Verify final count
        final_count = collection.count()
        logger.info(f"Final collection size: {final_count} documents")
        
        # Comprehensive testing
        logger.info("Performing comprehensive tests...")
        
        # Test 1: Numbered definition query
        test_queries = [
            "158. Diagnosis Code 6",
            "Diagnosis Code 6", 
            "HCFA Admit Type Code",
            "Technical Name: ICD9_DX_CD"
        ]
        
        for query in test_queries:
            try:
                test_embedding = embeddings.embed_query(query)
                test_results = collection.query(
                    query_embeddings=[test_embedding],
                    n_results=3,
                    include=["documents", "metadatas", "distances"]
                )
                
                if test_results and test_results.get('documents'):
                    logger.info(f"\nTest query: '{query}'")
                    logger.info(f"  Found {len(test_results['documents'][0])} results")
                    
                    first_result = test_results['documents'][0][0]
                    first_metadata = test_results.get('metadatas', [[]])[0][0] if test_results.get('metadatas') else {}
                    distance = test_results.get('distances', [[]])[0][0] if test_results.get('distances') else 1.0
                    
                    logger.info(f"  Top result distance: {distance:.4f}")
                    logger.info(f"  Content preview: {first_result[:100]}...")
                    
                    if first_metadata.get('is_numbered_definition'):
                        def_num = first_metadata.get('definition_number', 'N/A')
                        logger.info(f"  Found numbered definition: {def_num}")
                    
                    if first_metadata.get('is_embedded_table'):
                        table_type = first_metadata.get('table_type', 'unknown')
                        logger.info(f"  Found table: {table_type}")
                
            except Exception as e:
                logger.error(f"Test query '{query}' failed: {e}")
        
        # Save analysis report
        analysis_file = os.path.join(CHROMA_DB_DIR, "collection_analysis.json")
        with open(analysis_file, 'w') as f:
            json.dump(analysis, f, indent=2, default=str)
        logger.info(f"Saved collection analysis to {analysis_file}")
        
        logger.info("\n=== SUCCESS ===")
        logger.info("Comprehensive ChromaDB created successfully!")
        logger.info("Optimized for medical data dictionary with:")
        logger.info(f"  • {final_count} total documents")
        logger.info(f"  • {analysis['numbered_definitions']} numbered definitions")
        logger.info(f"  • {analysis['embedded_tables']} embedded tables/sheets")
        logger.info(f"  • {analysis['excel_sheets']} Excel sheets processed")
        logger.info(f"  • {analysis['complete_definitions']} complete definitions")
        logger.info(f"  • {analysis['quality_metrics']['definition_coverage']} unique definition numbers")
        logger.info("  • Comprehensive XLS/table processing")
        logger.info("  • Enhanced metadata for precise retrieval")
        
    except Exception as e:
        logger.error(f"Error creating ChromaDB: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if not os.path.exists(DOCUMENTS_DIR):
        logger.error(f"Documents directory '{DOCUMENTS_DIR}' not found")
        logger.info("Please create the directory and add your documents")
        os.makedirs(DOCUMENTS_DIR, exist_ok=True)
    elif not os.listdir(DOCUMENTS_DIR):
        logger.warning(f"Documents directory '{DOCUMENTS_DIR}' is empty")
        logger.info("Please add document files to this directory")
    else:
        main()