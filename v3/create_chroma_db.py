# Enhanced create_chroma_db.py for medical data dictionary
# Fixed for better numbered definition processing, enhanced XLS support, and detailed logging

import os
import shutil
import re
import time
from pathlib import Path
from langchain_community.document_loaders import (
    UnstructuredExcelLoader, 
    TextLoader,
    UnstructuredWordDocumentLoader,
    PyPDFLoader
)
from langchain_unstructured import UnstructuredLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_vertexai import VertexAIEmbeddings
from langchain.schema import Document
import chromadb
import logging
import json
from typing import List, Dict, Any, Optional, Tuple, Union
import pandas as pd
import numpy as np
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
import traceback

# Enhanced logging setup
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Configuration
DOCUMENTS_DIR = "./documents" 
CHROMA_DB_DIR = "./chroma_db_data"
COLLECTION_NAME = "rag_collection"
EMBEDDING_MODEL_NAME = "text-embedding-005"

# Enhanced chunking for complete numbered definitions and tables
CHUNK_SIZE = 3000  # Larger for complete definitions
CHUNK_OVERLAP = 500  # More overlap for context preservation
MIN_DEFINITION_SIZE = 50  # Minimum size for definition chunks
MAX_DEFINITION_SIZE = 2500  # Maximum size before forced splitting
TABLE_CHUNK_SIZE = 2000  # Special size for tables

# Initialize Vertex AI Embeddings
try:
    embeddings = VertexAIEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    logger.info(f"Successfully initialized VertexAIEmbeddings: {EMBEDDING_MODEL_NAME}")
except Exception as e:
    logger.error(f"Error initializing VertexAIEmbeddings: {e}")
    logger.info("Please ensure Google Cloud authentication and Vertex AI API access")
    exit(1)

def preprocess_text(text_content: str) -> str:
    """Enhanced text preprocessing for medical documentation"""
    if not text_content:
        return ""
    
    # Remove extra whitespace but preserve structure
    text_content = re.sub(r' +', ' ', text_content)
    text_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', text_content)
    
    # Clean non-printable characters but keep medical symbols and formatting
    text_content = re.sub(r'[^\x20-\x7E\n\r\t•\-–\—]', ' ', text_content)
    
    # Fix common OCR/extraction issues
    text_content = re.sub(r'(\d+)\s*\.\s*([A-Z])', r'\1. \2', text_content)  # Fix numbered definitions
    text_content = re.sub(r'([a-z])\s*:\s*([A-Z])', r'\1: \2', text_content)  # Fix field separators
    
    # Preserve bullet points and formatting
    text_content = re.sub(r'[•·]\s*', '• ', text_content)  # Normalize bullet points
    text_content = re.sub(r'\*\*([^*]+)\*\*', r'**\1**', text_content)  # Fix bold formatting
    
    return text_content.strip()

def detect_numbered_definition_boundaries(text: str) -> List[Tuple[int, int, str, str]]:
    """Enhanced detection of numbered definition boundaries with better pattern matching"""
    boundaries = []
    
    # Enhanced patterns for numbered definitions with more comprehensive coverage
    patterns = [
        r'(\d+)\.\s*([^:\n]+):\s*',                    # Standard: "158. Diagnosis Code 6:"
        r'\*\*(\d+)\.\s*([^:*\n]+):\*\*',             # Bold: "**158. Diagnosis Code 6:**"
        r'(\d+)\.\s*\[([^\]]+)\]\s*:',                # Bracketed: "158. [Diagnosis Code 6]:"
        r'(\d+)\.\s*([A-Z][^:\n]*[a-z])\s*:',         # Title case: "158. Diagnosis Code 6:"
        r'\*\*(\d+)\.\s*\[([^\]]+)\].*?:\*\*',        # Bold bracketed: "**158. [Diagnosis Code 6]:**"
        r'(\d+)\.\s*([A-Z][^:\n]*)',                  # Without colon: "158. Diagnosis Code 6"
    ]
    
    all_matches = []
    
    for pattern in patterns:
        matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
        for match in matches:
            start_pos = match.start()
            definition_number = match.group(1)
            definition_title = match.group(2).strip()
            # Clean title from markdown formatting
            definition_title = re.sub(r'[\[\]\*_{}]', '', definition_title).strip()
            full_match = match.group(0)
            all_matches.append((start_pos, definition_number, definition_title, full_match))
    
    # Sort by position and remove duplicates
    all_matches.sort(key=lambda x: x[0])
    
    # Remove overlapping matches
    final_matches = []
    last_end = -1
    
    for start_pos, number, title, full_match in all_matches:
        if start_pos > last_end:
            final_matches.append((start_pos, number, title, full_match))
            last_end = start_pos + len(full_match)
    
    # Convert to boundaries with end positions
    for i, (start_pos, number, title, full_match) in enumerate(final_matches):
        # Find the end of this definition
        if i + 1 < len(final_matches):
            end_pos = final_matches[i + 1][0]
        else:
            # Look for next major section or end of text
            next_section = re.search(r'\n\n[A-Z]', text[start_pos:])
            if next_section:
                end_pos = start_pos + next_section.start()
            else:
                end_pos = len(text)
        
        boundaries.append((start_pos, end_pos, f"{number}. {title}", number))
    
    logger.info(f"Detected {len(boundaries)} numbered definition boundaries")
    for start, end, title, num in boundaries[:5]:  # Log first 5
        logger.info(f"  {num}. {title} (pos: {start}-{end})")
    
    return boundaries

def extract_docx_tables(file_path: str) -> List[Dict]:
    """Enhanced table extraction from DOCX files with better error handling"""
    tables = []
    logger.info(f"Extracting tables from DOCX file: {file_path}")
    
    try:
        doc = DocxDocument(file_path)
        
        for table_idx, table in enumerate(doc.tables):
            logger.info(f"Processing table {table_idx + 1}/{len(doc.tables)}")
            table_data = []
            headers = []
            
            # Extract table headers
            if table.rows:
                header_row = table.rows[0]
                headers = []
                for cell in header_row.cells:
                    cell_text = cell.text.strip()
                    headers.append(cell_text)
                logger.info(f"  Table {table_idx + 1}: Found {len(headers)} headers")
                logger.debug(f"  Headers: {headers}")
            
            # Extract table data
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                if any(cell.strip() for cell in row_data):  # Skip empty rows
                    table_data.append(row_data)
            
            if table_data or headers:
                # Format table as text with enhanced structure
                formatted_content = []
                if headers:
                    formatted_content.append("Table Headers:")
                    formatted_content.append(" | ".join(headers))
                    formatted_content.append("-" * len(" | ".join(headers)))
                    formatted_content.append("")
                
                formatted_content.append("Table Data:")
                for row in table_data[:100]:  # Increased limit for better coverage
                    formatted_content.append(" | ".join(row))
                
                tables.append({
                    'type': 'docx_table',
                    'table_index': table_idx,
                    'headers': headers,
                    'data': table_data,
                    'formatted_content': '\n'.join(formatted_content),
                    'row_count': len(table_data),
                    'col_count': len(headers) if headers else 0,
                    'source_file': os.path.basename(file_path)
                })
                
                logger.info(f"  Successfully extracted DOCX table {table_idx + 1}: {len(table_data)} rows, {len(headers)} columns")
        
    except Exception as e:
        logger.error(f"Error extracting DOCX tables from {file_path}: {e}")
        logger.debug(f"Full traceback: {traceback.format_exc()}")
    
    logger.info(f"Total DOCX tables extracted: {len(tables)}")
    return tables

def extract_embedded_tables_advanced(content: str) -> List[Dict]:
    """Advanced table extraction with multiple detection methods and enhanced logging"""
    tables = []
    logger.debug("Starting advanced table extraction...")
    
    # Method 1: Enhanced value mapping tables
    value_patterns = [
        (r'([A-Z0-9]+)\s*[=\-]\s*([^\n]+)', 'equal_dash_mapping'),  # A = Description
        (r'([A-Z0-9]+)\s+([A-Z][^\n]+)', 'space_separated'),        # A Description
        (r'(\d+)\s+([A-Z][^\n]+)', 'number_description'),           # 1 Description
        (r'([A-Z0-9]{1,4})\s+([\w\s\-\']+?)(?=\n[A-Z0-9]{1,4}\s+|\n\n|\Z)', 'appendix_style'),
    ]
    
    for pattern, pattern_name in value_patterns:
        value_matches = re.findall(pattern, content, re.MULTILINE)
        if len(value_matches) > 3:  # Must have multiple entries
            tables.append({
                'type': 'value_mapping',
                'pattern_used': pattern_name,
                'data': value_matches,
                'formatted_content': '\n'.join([f"{code} = {desc}" for code, desc in value_matches]),
                'entry_count': len(value_matches)
            })
            logger.info(f"Found value mapping table using {pattern_name}: {len(value_matches)} entries")
            break
    
    # Method 2: Enhanced field specification tables
    field_specs = re.findall(r'(Format|Technical Name|Length|Position[s]?):\s*([^\n]+)', content)
    if field_specs:
        tables.append({
            'type': 'field_specification',
            'data': field_specs,
            'formatted_content': '\n'.join([f"• **{field}:** {value}" for field, value in field_specs]),
            'field_count': len(field_specs)
        })
        logger.info(f"Found field specification table: {len(field_specs)} fields")
    
    # Method 3: Enhanced appendix/lookup tables with headers
    lines = content.split('\n')
    table_blocks = []
    current_block = []
    in_table = False
    
    for line in lines:
        # Enhanced table detection patterns
        if (re.match(r'^[A-Z][^|]*\|[^|]*', line) or 
            re.match(r'^[A-Z][a-z]+\s+[A-Z][a-z]+', line) or
            '---' in line or '===' in line or
            re.match(r'^[A-Z0-9]{1,4}\s+[A-Z]', line)):
            if not in_table:
                in_table = True
                current_block = [line]
            else:
                current_block.append(line)
        elif in_table:
            if line.strip() and not re.match(r'^\d+\.', line):  # Continue table
                current_block.append(line)
            else:  # End of table
                if len(current_block) > 2:
                    table_blocks.append(current_block)
                current_block = []
                in_table = False
    
    # Process detected table blocks with enhanced logic
    for i, block in enumerate(table_blocks):
        cleaned_block = [line.strip() for line in block if line.strip()]
        if len(cleaned_block) > 2:
            # Try to detect headers
            headers = []
            data_rows = cleaned_block
            
            # Check if first row looks like headers
            first_row = cleaned_block[0]
            if ('|' in first_row or 
                re.match(r'^[A-Z][^a-z]*\s+[A-Z][^a-z]*', first_row)):
                headers = re.split(r'\s*\|\s*', first_row) if '|' in first_row else first_row.split()
                data_rows = cleaned_block[1:]
            
            tables.append({
                'type': 'structured_table',
                'table_index': i,
                'headers': headers,
                'data_rows': data_rows,
                'formatted_content': '\n'.join(cleaned_block),
                'row_count': len(data_rows)
            })
            logger.info(f"Found structured table {i}: {len(data_rows)} rows")
    
    # Method 4: Enhanced code tables from appendices
    appendix_pattern = r'([A-Z0-9]{1,4})\s+([\w\s\-\']+?)(?=\n[A-Z0-9]{1,4}\s+|\n\n|\Z)'
    appendix_matches = re.findall(appendix_pattern, content, re.MULTILINE)
    
    if len(appendix_matches) > 5:  # Substantial lookup table
        tables.append({
            'type': 'appendix_lookup',
            'data': appendix_matches,
            'formatted_content': '\n'.join([f"{code:5s} {desc}" for code, desc in appendix_matches]),
            'entry_count': len(appendix_matches)
        })
        logger.info(f"Found appendix lookup table: {len(appendix_matches)} entries")
    
    logger.info(f"Total embedded tables extracted: {len(tables)}")
    return tables

def analyze_content_patterns_advanced(text: str) -> Dict[str, Any]:
    """Enhanced content pattern analysis with better numbered definition detection"""
    patterns = {}
    logger.debug("Starting advanced content pattern analysis...")
    
    # 1. Enhanced numbered definitions detection
    numbered_def_patterns = [
        (r'(\d+)\.\s*([^:\n]+):', 'standard_colon'),
        (r'\*\*(\d+)\.\s*([^:*\n]+):\*\*', 'bold_colon'),
        (r'(\d+)\.\s*\[([^\]]+)\]:', 'bracketed_colon'),
        (r'\*\*(\d+)\.\s*\[([^\]]+)\].*?:\*\*', 'bold_bracketed'),
        (r'(\d+)\.\s*([A-Z][^:\n]*)', 'no_colon'),
    ]
    
    all_numbered_defs = []
    pattern_counts = {}
    
    for pattern, pattern_name in numbered_def_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
        pattern_counts[pattern_name] = len(matches)
        all_numbered_defs.extend(matches)
    
    logger.info(f"Numbered definition patterns found: {pattern_counts}")
    
    # Remove duplicates and sort
    unique_defs = []
    seen = set()
    for match in all_numbered_defs:
        key = match[0]  # Use the number as key
        if key not in seen:
            seen.add(key)
            unique_defs.append(match)
    
    patterns['numbered_definitions'] = unique_defs
    patterns['definition_count'] = len(unique_defs)
    
    # Extract definition numbers for reference
    definition_numbers = sorted(set([match[0] for match in unique_defs if match[0].isdigit()]))
    patterns['definition_numbers'] = definition_numbers
    
    logger.info(f"Found {len(unique_defs)} unique numbered definitions")
    logger.info(f"Definition numbers: {definition_numbers[:10]}{'...' if len(definition_numbers) > 10 else ''}")
    
    # 2. Enhanced technical names detection
    tech_name_patterns = [
        r'Technical Name:\s*([A-Z_][A-Z0-9_]*)',
        r'Technical Name:\s*([A-Z][A-Z0-9\s_]*)',
        r'\*\*Technical Name:\*\*\s*([A-Z_][A-Z0-9_]*)',
    ]
    
    technical_names = []
    for pattern in tech_name_patterns:
        matches = re.findall(pattern, text)
        technical_names.extend(matches)
    
    patterns['technical_names'] = list(set(technical_names))
    logger.info(f"Found {len(patterns['technical_names'])} technical names")
    
    # 3. Enhanced field specifications 
    field_elements = {
        'formats': re.findall(r'[*•]\s*\*\*Format:\*\*\s*([^\n]+)', text),
        'lengths': re.findall(r'[*•]\s*\*\*Length:\*\*\s*(\d+(?:\.\d+)?)', text),
        'positions': re.findall(r'[*•]\s*\*\*Position[s]?:\*\*\s*([\d\s\-–—,]+)', text),
        'definitions': re.findall(r'[*•]\s*\*\*Definition:\*\*\s*([^\n]+)', text)
    }
    
    patterns.update(field_elements)
    logger.info(f"Field elements found: {[(k, len(v)) for k, v in field_elements.items()]}")
    
    # 4. Enhanced value mappings and codes
    value_mappings = []
    value_patterns = [
        r'([A-Z0-9]{1,4})\s*=\s*([^\n]+)',
        r'([A-Z0-9]{1,4})\s+([A-Z][^\n]+)',
        r'(\d+)\s+([A-Z][A-Za-z\s]+?)(?=\n\d+|\n\n|\Z)'
    ]
    
    for pattern in value_patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        if len(matches) > 3:  # Substantial list
            value_mappings.extend(matches)
    
    patterns['value_mappings'] = value_mappings[:100]  # Limit for performance
    logger.info(f"Found {len(value_mappings)} value mappings")
    
    # 5. Cross-references and relationships
    field_refs = re.findall(r'[Ff]ield\s*#?\s*(\d+)', text)
    appendix_refs = len(re.findall(r'[Aa]ppendix|[Aa]ppendices', text))
    see_also_refs = re.findall(r'[Ss]ee\s+(?:also\s+)?[Ff]ield\s*#?(\d+)', text)
    
    patterns['field_references'] = field_refs
    patterns['see_also_references'] = see_also_refs
    patterns['has_appendix_refs'] = appendix_refs > 0
    patterns['appendix_ref_count'] = appendix_refs
    
    logger.info(f"References found - Field: {len(field_refs)}, See also: {len(see_also_refs)}, Appendix: {appendix_refs}")
    
    # 6. Enhanced content classification
    content_type = 'descriptive_text'  # Default
    
    if unique_defs:
        if len(unique_defs) == 1:
            content_type = 'single_numbered_definition'
        else:
            content_type = 'multiple_numbered_definitions'
    elif technical_names and any(field_elements.values()):
        content_type = 'field_specification'
    elif len(value_mappings) > 10:
        content_type = 'value_mapping_table'
    elif appendix_refs > 0 and len(value_mappings) > 3:
        content_type = 'appendix_table'
    elif 'Aetna' in text and ('Field' in text or 'Definition' in text):
        content_type = 'medical_data_dictionary'
    
    patterns['content_type'] = content_type
    
    # 7. Enhanced quality indicators
    patterns['has_complete_structure'] = (
        len(field_elements['formats']) > 0 and
        len(patterns['technical_names']) > 0 and
        len(field_elements['definitions']) > 0
    )
    
    patterns['completeness_score'] = sum([
        1 if field_elements['formats'] else 0,
        1 if patterns['technical_names'] else 0,
        1 if field_elements['lengths'] else 0,
        1 if field_elements['positions'] else 0,
        1 if field_elements['definitions'] else 0
    ])
    
    # 8. Table indicators
    patterns['has_tabular_data'] = (
        '|' in text or 
        len(value_mappings) > 5 or
        re.search(r'\n[A-Z0-9]+\s+[A-Z]', text) is not None
    )
    
    logger.info(f"Content classified as: {content_type} (completeness: {patterns['completeness_score']}/5)")
    
    return patterns

def create_enhanced_metadata(doc: Document, patterns: Dict, tables: List[Dict], file_path: str, chunk_info: Dict = None) -> Dict:
    """Create comprehensive metadata for ChromaDB with enhanced field detection"""
    metadata = {
        'source_filename': os.path.basename(file_path),
        'content_length': len(doc.page_content),
        'content_type': patterns.get('content_type', 'unknown'),
    }
    
    # Add original document metadata
    if hasattr(doc, 'metadata') and doc.metadata:
        metadata.update({
            'page_number': str(doc.metadata.get('page_number', 'N/A')),
            'element_type': doc.metadata.get('category', doc.metadata.get('element_type', 'Unknown')),
        })
    else:
        metadata.update({
            'page_number': 'N/A',
            'element_type': 'Unknown',
        })
    
    # Chunk information
    if chunk_info:
        metadata.update(chunk_info)
    
    # Enhanced numbered definition metadata 
    numbered_defs = patterns.get('numbered_definitions', [])
    if numbered_defs:
        # Handle both single and multiple definitions
        definition_numbers = [def_info[0] for def_info in numbered_defs if len(def_info) > 0]
        definition_titles = [def_info[1] for def_info in numbered_defs if len(def_info) > 1]
        
        metadata['definition_numbers'] = ','.join(definition_numbers)
        metadata['definition_titles'] = ','.join(definition_titles)
        metadata['is_numbered_definition'] = True
        metadata['definition_count'] = len(numbered_defs)
        
        # Single definition metadata
        if len(numbered_defs) == 1:
            metadata['definition_number'] = definition_numbers[0]
            metadata['definition_title'] = definition_titles[0] if definition_titles else ''
        
        # Check for specific fields with enhanced patterns
        for number, title in numbered_defs:
            title_lower = title.lower()
            if 'diagnosis code' in title_lower:
                metadata['is_diagnosis_field'] = True
                metadata['diagnosis_type'] = title
            elif 'hcfa' in title_lower:
                metadata['is_hcfa_field'] = True
                if 'admit type' in title_lower:
                    metadata['hcfa_field_type'] = 'admit_type'
                elif 'admit source' in title_lower:
                    metadata['hcfa_field_type'] = 'admit_source'
                elif 'place of service' in title_lower:
                    metadata['hcfa_field_type'] = 'place_of_service'
            elif 'admit' in title_lower:
                metadata['is_admit_field'] = True
    else:
        metadata['is_numbered_definition'] = False
        metadata['definition_count'] = 0
    
    # Enhanced technical name metadata
    technical_names = patterns.get('technical_names', [])
    if technical_names:
        metadata['technical_names'] = ','.join(technical_names)
        metadata['is_field_spec'] = True
        metadata['technical_name_count'] = len(technical_names)
        
        if len(technical_names) == 1:
            metadata['primary_technical_name'] = technical_names[0]
            
            # Enhanced classification of technical name types
            tech_name = technical_names[0]
            if 'DX_CD' in tech_name:
                metadata['field_category'] = 'diagnosis'
            elif 'AMT' in tech_name or 'AMOUNT' in tech_name:
                metadata['field_category'] = 'amount'
            elif '_CD' in tech_name:
                metadata['field_category'] = 'code'
            elif '_DT' in tech_name or 'DATE' in tech_name:
                metadata['field_category'] = 'date'
            elif '_ID' in tech_name:
                metadata['field_category'] = 'identifier'
            elif 'HCFA' in tech_name:
                metadata['field_category'] = 'hcfa'
    else:
        metadata['is_field_spec'] = False
        metadata['technical_name_count'] = 0
    
    # Enhanced field specification elements
    field_elements = ['formats', 'lengths', 'positions', 'definitions']
    for element in field_elements:
        values = patterns.get(element, [])
        if values:
            metadata[f'{element[:-1]}_info'] = ','.join(str(v) for v in values[:3])  # Limit for metadata size
    
    # Content quality and completeness
    completeness_score = patterns.get('completeness_score', 0)
    metadata['definition_completeness'] = completeness_score
    metadata['is_complete_definition'] = completeness_score >= 4
    metadata['has_complete_structure'] = patterns.get('has_complete_structure', False)
    
    # Enhanced table and structured data metadata
    metadata['has_embedded_tables'] = len(tables) > 0
    metadata['table_count'] = len(tables)
    metadata['has_tabular_data'] = patterns.get('has_tabular_data', False)
    
    if tables:
        table_types = list(set([table['type'] for table in tables]))
        metadata['table_types'] = ','.join(table_types)
        
        # Specific table information
        total_entries = sum(table.get('entry_count', table.get('row_count', 0)) for table in tables)
        metadata['total_table_entries'] = total_entries
        
        # Detailed table analysis
        docx_tables = [t for t in tables if t['type'] == 'docx_table']
        value_mappings = [t for t in tables if t['type'] == 'value_mapping']
        
        metadata['docx_table_count'] = len(docx_tables)
        metadata['value_mapping_count'] = len(value_mappings)
    
    # Enhanced value mapping metadata
    value_mappings = patterns.get('value_mappings', [])
    metadata['has_value_mappings'] = len(value_mappings) > 0
    metadata['value_mapping_count'] = len(value_mappings)
    
    # Reference and relationship metadata
    field_refs = patterns.get('field_references', [])
    see_also_refs = patterns.get('see_also_references', [])
    
    metadata['field_references'] = ','.join(field_refs) if field_refs else ''
    metadata['see_also_references'] = ','.join(see_also_refs) if see_also_refs else ''
    metadata['has_field_references'] = len(field_refs) > 0
    metadata['has_see_also_references'] = len(see_also_refs) > 0
    metadata['has_appendix_references'] = patterns.get('has_appendix_refs', False)
    metadata['appendix_ref_count'] = patterns.get('appendix_ref_count', 0)
    
    # Enhanced classification indicators
    metadata['is_appendix'] = patterns.get('content_type') in ['appendix_table', 'value_mapping_table']
    metadata['is_lookup_table'] = patterns.get('content_type') == 'value_mapping_table'
    metadata['is_single_definition'] = patterns.get('content_type') == 'single_numbered_definition'
    metadata['is_multiple_definitions'] = patterns.get('content_type') == 'multiple_numbered_definitions'
    
    # Enhanced search optimization metadata
    searchable_terms = []
    
    # Add numbered definition terms with variations
    if numbered_defs:
        for number, title in numbered_defs:
            searchable_terms.extend([
                f"{number}. {title}",
                f"**{number}. {title}:**",
                f"{number}. [{title}]",
                title.upper(),
                title.lower(),
                title.replace(' ', '_').upper(),  # Technical name format
                number  # Just the number
            ])
    
    # Add technical names with variations
    if technical_names:
        searchable_terms.extend(technical_names)
        searchable_terms.extend([name.lower() for name in technical_names])
        searchable_terms.extend([name.replace('_', ' ') for name in technical_names])
    
    # Add field references
    if field_refs:
        searchable_terms.extend([f"Field {ref}" for ref in field_refs])
        searchable_terms.extend([f"Field #{ref}" for ref in field_refs])
    
    # Limit searchable terms for ChromaDB constraints
    if searchable_terms:
        unique_terms = list(dict.fromkeys(searchable_terms))  # Remove duplicates while preserving order
        metadata['searchable_terms'] = ','.join(unique_terms[:25])  # Increased limit
    
    # Enhanced content categorization for better retrieval
    content_categories = []
    if metadata.get('is_numbered_definition'):
        content_categories.append('numbered_definition')
    if metadata.get('is_field_spec'):
        content_categories.append('field_specification')
    if metadata.get('has_embedded_tables'):
        content_categories.append('contains_tables')
    if metadata.get('is_appendix'):
        content_categories.append('appendix')
    if metadata.get('has_value_mappings'):
        content_categories.append('value_mappings')
    if metadata.get('is_hcfa_field'):
        content_categories.append('hcfa_field')
    
    metadata['content_categories'] = ','.join(content_categories) if content_categories else ''
    
    # Ensure all metadata values are ChromaDB compatible
    for key, value in metadata.items():
        if isinstance(value, (list, tuple)):
            metadata[key] = ','.join(str(v) for v in value)
        elif not isinstance(value, (str, int, float, bool, type(None))):
            metadata[key] = str(value)
        elif value is None:
            metadata[key] = ''
        # Handle very long strings
        elif isinstance(value, str) and len(value) > 500:
            metadata[key] = value[:500] + '...'
    
    return metadata

def smart_chunk_numbered_definitions(text: str, text_splitter: RecursiveCharacterTextSplitter) -> List[Document]:
    """Enhanced chunking that preserves complete numbered definitions"""
    boundaries = detect_numbered_definition_boundaries(text)
    chunks = []
    
    if not boundaries:
        # No numbered definitions found, use standard chunking
        return text_splitter.create_documents([text])
    
    current_pos = 0
    
    for start_pos, end_pos, definition_title, definition_number in boundaries:
        # Add any text before the definition
        if current_pos < start_pos:
            pre_text = text[current_pos:start_pos].strip()
            if pre_text and len(pre_text) > MIN_DEFINITION_SIZE:
                pre_chunks = text_splitter.create_documents([pre_text])
                chunks.extend(pre_chunks)
        
        # Extract the complete definition
        definition_text = text[start_pos:end_pos].strip()
        definition_title_clean = definition_title.replace(f"{definition_number}. ", "")
        
        if len(definition_text) <= MAX_DEFINITION_SIZE:
            # Keep complete definition as one chunk
            chunk_doc = Document(
                page_content=definition_text,
                metadata={
                    'chunk_type': 'complete_definition',
                    'definition_number': definition_number,
                    'definition_title': definition_title_clean,
                    'start_pos': start_pos,
                    'end_pos': end_pos,
                    'definition_length': len(definition_text)
                }
            )
            chunks.append(chunk_doc)
            logger.info(f"Created complete definition chunk: {definition_number}. {definition_title_clean}")
        else:
            # Definition is too long, split carefully at logical boundaries
            lines = definition_text.split('\n')
            current_chunk_lines = []
            current_length = 0
            header_lines = []
            
            # Extract header (first few lines with definition title)
            for i, line in enumerate(lines[:5]):
                if definition_number in line or definition_title_clean.split()[-1] in line:
                    header_lines.append(line)
                else:
                    break
            
            chunk_index = 0
            for line in lines:
                # Check if adding this line exceeds size limit
                if current_length + len(line) > CHUNK_SIZE and current_chunk_lines:
                    # Create chunk with header if it's not the first chunk
                    chunk_content = '\n'.join(current_chunk_lines)
                    if chunk_index > 0 and header_lines:
                        chunk_content = '\n'.join(header_lines) + '\n...\n' + chunk_content
                    
                    chunk_doc = Document(
                        page_content=chunk_content,
                        metadata={
                            'chunk_type': 'partial_definition',
                            'definition_number': definition_number,
                            'definition_title': definition_title_clean,
                            'chunk_index': chunk_index,
                            'is_continuation': chunk_index > 0,
                            'has_header': chunk_index > 0 and bool(header_lines)
                        }
                    )
                    chunks.append(chunk_doc)
                    logger.info(f"Created partial definition chunk: {definition_number}. {definition_title_clean} (part {chunk_index + 1})")
                    
                    current_chunk_lines = []
                    current_length = 0
                    chunk_index += 1
                
                current_chunk_lines.append(line)
                current_length += len(line) + 1
            
            # Add final chunk
            if current_chunk_lines:
                chunk_content = '\n'.join(current_chunk_lines)
                if chunk_index > 0 and header_lines:
                    chunk_content = '\n'.join(header_lines) + '\n...\n' + chunk_content
                
                chunk_doc = Document(
                    page_content=chunk_content,
                    metadata={
                        'chunk_type': 'partial_definition',
                        'definition_number': definition_number,
                        'definition_title': definition_title_clean,
                        'chunk_index': chunk_index,
                        'is_continuation': chunk_index > 0,
                        'is_final_chunk': True,
                        'has_header': chunk_index > 0 and bool(header_lines)
                    }
                )
                chunks.append(chunk_doc)
                logger.info(f"Created final partial definition chunk: {definition_number}. {definition_title_clean} (part {chunk_index + 1})")
        
        current_pos = end_pos
    
    # Add any remaining text
    if current_pos < len(text):
        remaining_text = text[current_pos:].strip()
        if remaining_text and len(remaining_text) > MIN_DEFINITION_SIZE:
            remaining_chunks = text_splitter.create_documents([remaining_text])
            chunks.extend(remaining_chunks)
            logger.info(f"Added {len(remaining_chunks)} chunks from remaining text")
    
    logger.info(f"Smart chunking completed: {len(chunks)} total chunks")
    return chunks

def load_document_with_comprehensive_extraction(file_path: str) -> List[Document]:
    """Enhanced document loading with comprehensive extraction and better error handling"""
    extension = os.path.splitext(file_path)[1].lower()
    documents = []
    
    logger.info(f"Loading document: {file_path} ({extension})")
    
    # Choose appropriate loader and extract method
    try:
        if extension == ".docx":
            # Use both unstructured loader and direct DOCX processing
            logger.info("Using DOCX loader with element mode...")
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            raw_docs = loader.load()
            
            # Also extract tables directly from DOCX
            logger.info("Extracting tables from DOCX...")
            docx_tables = extract_docx_tables(file_path)
            
        elif extension == ".doc":
            logger.info("Using DOC loader...")
            loader = UnstructuredLoader(file_path, mode="elements", strategy="auto")
            raw_docs = loader.load()
            docx_tables = []
            
        elif extension in [".xlsx", ".xls"]:
            # Handle Excel files specially
            logger.info("Processing Excel file...")
            return load_excel_comprehensive(file_path)
            
        elif extension == ".pdf":
            logger.info("Using PDF loader...")
            loader = PyPDFLoader(file_path)
            raw_docs = loader.load()
            docx_tables = []
            
        elif extension == ".txt":
            logger.info("Using text loader...")
            loader = TextLoader(file_path, encoding="utf-8")
            raw_docs = loader.load()
            docx_tables = []
            
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return documents
        
        logger.info(f"Loaded {len(raw_docs)} raw documents")
        
        # Process main content
        combined_content = ""
        for doc in raw_docs:
            doc.page_content = preprocess_text(doc.page_content)
            if len(doc.page_content.strip()) >= MIN_DEFINITION_SIZE:
                combined_content += doc.page_content + "\n\n"
        
        logger.info(f"Combined content length: {len(combined_content)} characters")
        
        # Analyze patterns in combined content
        logger.info("Analyzing content patterns...")
        patterns = analyze_content_patterns_advanced(combined_content)
        
        logger.info("Extracting embedded tables...")
        embedded_tables = extract_embedded_tables_advanced(combined_content)
        
        # Combine extracted tables
        all_tables = embedded_tables + docx_tables
        logger.info(f"Total tables found: {len(embedded_tables)} embedded + {len(docx_tables)} DOCX = {len(all_tables)}")
        
        # Create document with combined content
        if combined_content.strip():
            main_doc = Document(
                page_content=combined_content.strip(),
                metadata={
                    'source_filename': os.path.basename(file_path),
                    'element_type': 'combined_content',
                    'page_number': 'combined',
                }
            )
            
            # Create enhanced metadata
            enhanced_metadata = create_enhanced_metadata(main_doc, patterns, all_tables, file_path)
            main_doc.metadata = enhanced_metadata
            
            documents.append(main_doc)
            logger.info("Created main document with enhanced metadata")
        
        # Create separate documents for each extracted table
        for i, table in enumerate(all_tables):
            table_doc = Document(
                page_content=table['formatted_content'],
                metadata={
                    'source_filename': os.path.basename(file_path),
                    'element_type': f"table_{table['type']}",
                    'table_index': i,
                    'table_type': table['type'],
                    'is_embedded_table': True,
                    'is_table_content': True,
                    'table_entry_count': table.get('entry_count', table.get('row_count', 0)),
                    'page_number': 'table',
                }
            )
            
            # Add table-specific metadata
            table_metadata = create_enhanced_metadata(table_doc, {
                'content_type': 'table_content',
                'has_tabular_data': True,
                'value_mappings': table.get('data', [])
            }, [], file_path)
            table_doc.metadata.update(table_metadata)
            
            documents.append(table_doc)
            logger.info(f"Created table document {i}: {table['type']}")
    
    except Exception as e:
        logger.error(f"Error loading {file_path}: {e}")
        logger.debug(f"Full traceback: {traceback.format_exc()}")
        
        # Try fallback loading for critical files
        try:
            logger.info("Attempting fallback loading...")
            if extension in [".docx", ".doc"]:
                loader = TextLoader(file_path, encoding="utf-8")
                fallback_docs = loader.load()
                for doc in fallback_docs:
                    doc.page_content = preprocess_text(doc.page_content)
                    doc.metadata['source_filename'] = os.path.basename(file_path)
                    doc.metadata['element_type'] = 'fallback_text'
                    documents.append(doc)
                logger.info(f"Fallback loading successful: {len(fallback_docs)} documents")
        except Exception as fallback_error:
            logger.error(f"Fallback loading also failed: {fallback_error}")
    
    logger.info(f"Document loading complete: {len(documents)} documents created")
    return documents

def load_excel_comprehensive(file_path: str) -> List[Document]:
    """Enhanced Excel file processing with detailed logging and error handling"""
    documents = []
    logger.info(f"Starting comprehensive Excel processing for: {file_path}")
    
    processing_stats = {
        'total_sheets': 0,
        'successful_sheets': 0,
        'failed_sheets': 0,
        'total_rows': 0,
        'total_columns': 0,
        'lookup_tables_found': 0,
        'value_mappings_found': 0
    }
    
    try:
        # Method 1: Use pandas for structured data
        excel_file = pd.ExcelFile(file_path)
        processing_stats['total_sheets'] = len(excel_file.sheet_names)
        logger.info(f"Found {processing_stats['total_sheets']} sheets in Excel file")
        
        for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
            logger.info(f"Processing sheet {sheet_idx + 1}/{processing_stats['total_sheets']}: '{sheet_name}'")
            
            try:
                # Read sheet with various options to handle different formats
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=0)
                
                if df.empty:
                    logger.warning(f"  Sheet '{sheet_name}' is empty, skipping")
                    continue
                
                # Clean up the dataframe
                df = df.dropna(how='all')  # Remove completely empty rows
                df = df.loc[:, ~df.columns.str.contains('^Unnamed')]  # Remove unnamed columns
                
                logger.info(f"  Sheet dimensions: {len(df)} rows x {len(df.columns)} columns")
                processing_stats['total_rows'] += len(df)
                processing_stats['total_columns'] += len(df.columns)
                
                # Convert DataFrame to structured text
                content_lines = [f"**Sheet: {sheet_name}**", ""]
                
                # Add metadata about the sheet
                content_lines.append(f"**Sheet Info:**")
                content_lines.append(f"• Rows: {len(df)}")
                content_lines.append(f"• Columns: {len(df.columns)}")
                content_lines.append("")
                
                # Add column headers
                headers = df.columns.tolist()
                content_lines.append("**Column Headers:**")
                for i, col in enumerate(headers):
                    content_lines.append(f"{i+1}. {col}")
                content_lines.append("")
                
                # Add data in tabular format (limited rows for performance)
                content_lines.append("**Data Preview:**")
                content_lines.append(" | ".join(str(col) for col in headers))
                content_lines.append("-" * (sum(len(str(col)) + 3 for col in headers)))
                
                for idx, row in df.head(200).iterrows():  # Increased preview limit
                    row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
                    content_lines.append(row_text)
                
                if len(df) > 200:
                    content_lines.append(f"... and {len(df) - 200} more rows")
                
                # Analyze the sheet content
                sheet_content = "\n".join(content_lines)
                
                # Enhanced analysis to determine if this is a lookup table
                is_lookup = False
                has_codes = False
                is_value_mapping = False
                
                if len(df) > 5 and len(df.columns) >= 2:
                    # Check first column for codes/values
                    first_col = df.iloc[:, 0].astype(str)
                    second_col = df.iloc[:, 1].astype(str)
                    
                    # Pattern 1: Short codes in first column
                    short_codes = first_col.str.match(r'^[A-Z0-9]{1,10}$')
                    if short_codes.sum() > len(df) * 0.7:  # 70% are short codes
                        has_codes = True
                        is_lookup = True
                        processing_stats['lookup_tables_found'] += 1
                        logger.info(f"    Detected lookup table with {short_codes.sum()} codes")
                    
                    # Pattern 2: Value mapping (code = description)
                    if len(df.columns) >= 2:
                        # Check if second column has descriptions
                        avg_desc_length = second_col.str.len().mean()
                        if avg_desc_length > 10 and has_codes:  # Non-empty descriptions
                            is_value_mapping = True
                            processing_stats['value_mappings_found'] += 1
                            logger.info(f"    Detected value mapping table")
                    
                    # Pattern 3: Numerical data tables
                    numeric_cols = df.select_dtypes(include=[np.number]).columns
                    if len(numeric_cols) > len(df.columns) * 0.5:
                        logger.info(f"    Detected numerical data table with {len(numeric_cols)} numeric columns")
                
                # Create document with enhanced metadata
                sheet_doc = Document(
                    page_content=sheet_content,
                    metadata={
                        'source_filename': os.path.basename(file_path),
                        'sheet_name': sheet_name,
                        'sheet_index': sheet_idx,
                        'element_type': 'excel_sheet',
                        'is_embedded_table': True,
                        'is_table_content': True,
                        'table_type': 'excel_data',
                        'row_count': len(df),
                        'column_count': len(df.columns),
                        'has_headers': True,
                        'is_lookup_table': is_lookup,
                        'is_value_mapping': is_value_mapping,
                        'has_codes': has_codes,
                        'content_type': 'value_mapping_table' if is_value_mapping else 'data_table',
                        'column_headers': ','.join(headers[:10])  # Store first 10 headers
                    }
                )
                
                # Add patterns analysis
                patterns = {
                    'content_type': 'value_mapping_table' if is_value_mapping else 'data_table',
                    'has_tabular_data': True,
                    'value_mappings': []
                }
                
                if is_value_mapping and len(df.columns) >= 2:
                    # Extract value mappings for pattern analysis
                    for _, row in df.head(100).iterrows():  # Analyze first 100 rows
                        code = str(row.iloc[0]) if pd.notna(row.iloc[0]) else ""
                        desc = str(row.iloc[1]) if pd.notna(row.iloc[1]) else ""
                        if code and desc and len(code) <= 10:  # Valid mapping
                            patterns['value_mappings'].append((code, desc))
                
                # Create enhanced metadata
                enhanced_metadata = create_enhanced_metadata(sheet_doc, patterns, [], file_path)
                sheet_doc.metadata.update(enhanced_metadata)
                
                documents.append(sheet_doc)
                processing_stats['successful_sheets'] += 1
                
                logger.info(f"  ✓ Successfully processed sheet '{sheet_name}'")
                
            except Exception as e:
                processing_stats['failed_sheets'] += 1
                logger.warning(f"  ✗ Could not process sheet '{sheet_name}': {e}")
                logger.debug(f"  Error details: {traceback.format_exc()}")
    
    except Exception as e:
        logger.error(f"Error processing Excel file {file_path}: {e}")
        logger.debug(f"Full traceback: {traceback.format_exc()}")
        
        # Fallback to unstructured loader
        try:
            logger.info("Attempting fallback Excel loading with UnstructuredExcelLoader...")
            loader = UnstructuredExcelLoader(file_path, mode="elements")
            fallback_docs = loader.load()
            for doc in fallback_docs:
                doc.page_content = preprocess_text(doc.page_content)
                doc.metadata['source_filename'] = os.path.basename(file_path)
                doc.metadata['element_type'] = 'excel_fallback'
                doc.metadata['is_table_content'] = True
                documents.append(doc)
            logger.info(f"Fallback Excel loading successful: {len(fallback_docs)} documents")
            processing_stats['successful_sheets'] = len(fallback_docs)
        except Exception as fallback_error:
            logger.error(f"Fallback Excel loading also failed: {fallback_error}")
    
    # Final logging summary with detailed statistics
    logger.info("=== EXCEL PROCESSING SUMMARY ===")
    logger.info(f"File: {os.path.basename(file_path)}")
    logger.info(f"Sheets found: {processing_stats['total_sheets']}")
    logger.info(f"Sheets successfully processed: {processing_stats['successful_sheets']}")
    logger.info(f"Sheets failed: {processing_stats['failed_sheets']}")
    logger.info(f"Total rows processed: {processing_stats['total_rows']}")
    logger.info(f"Total columns processed: {processing_stats['total_columns']}")
    logger.info(f"Lookup tables found: {processing_stats['lookup_tables_found']}")
    logger.info(f"Value mapping tables found: {processing_stats['value_mappings_found']}")
    logger.info(f"Documents created: {len(documents)}")
    logger.info("=================================")
    
    return documents

def load_all_documents(source_dir: str) -> List[Document]:
    """Enhanced document loading with comprehensive logging and statistics"""
    all_documents = []
    processing_stats = {
        'total_files_found': 0,
        'files_processed': 0,
        'files_failed': 0,
        'excel_files': 0,
        'excel_sheets': 0,
        'docx_files': 0,
        'pdf_files': 0,
        'txt_files': 0,
        'total_documents': 0,
        'numbered_definitions': 0,
        'tables_extracted': 0,
        'complete_definitions': 0
    }
    
    if not os.path.exists(source_dir):
        logger.error(f"Source directory '{source_dir}' not found")
        return all_documents
    
    supported_extensions = {'.docx', '.doc', '.xlsx', '.xls', '.pdf', '.txt'}
    
    # Scan directory first
    files_to_process = []
    for filename in os.listdir(source_dir):
        if filename.startswith('.'):
            continue
        
        file_path = os.path.join(source_dir, filename)
        extension = os.path.splitext(filename)[1].lower()
        
        if extension in supported_extensions and os.path.isfile(file_path):
            files_to_process.append((file_path, extension))
            processing_stats['total_files_found'] += 1
            
            # Count by file type
            if extension in ['.xlsx', '.xls']:
                processing_stats['excel_files'] += 1
            elif extension == '.docx':
                processing_stats['docx_files'] += 1
            elif extension == '.pdf':
                processing_stats['pdf_files'] += 1
            elif extension == '.txt':
                processing_stats['txt_files'] += 1
    
    logger.info(f"Found {processing_stats['total_files_found']} files to process:")
    logger.info(f"  Excel files: {processing_stats['excel_files']}")
    logger.info(f"  DOCX files: {processing_stats['docx_files']}")
    logger.info(f"  PDF files: {processing_stats['pdf_files']}")
    logger.info(f"  Text files: {processing_stats['txt_files']}")
    
    # Process each file
    for file_path, extension in files_to_process:
        filename = os.path.basename(file_path)
        logger.info(f"\n--- Processing file {processing_stats['files_processed'] + 1}/{processing_stats['total_files_found']}: {filename} ---")
        
        try:
            docs = load_document_with_comprehensive_extraction(file_path)
            all_documents.extend(docs)
            processing_stats['files_processed'] += 1
            processing_stats['total_documents'] += len(docs)
            
            # Analyze the extracted documents
            numbered_defs = sum(1 for doc in docs if doc.metadata.get('is_numbered_definition'))
            tables = sum(1 for doc in docs if doc.metadata.get('is_embedded_table'))
            complete_defs = sum(1 for doc in docs if doc.metadata.get('is_complete_definition'))
            
            processing_stats['numbered_definitions'] += numbered_defs
            processing_stats['tables_extracted'] += tables
            processing_stats['complete_definitions'] += complete_defs
            
            # Track Excel-specific stats
            if extension in ['.xlsx', '.xls']:
                excel_sheets = sum(1 for doc in docs if doc.metadata.get('element_type') == 'excel_sheet')
                processing_stats['excel_sheets'] += excel_sheets
                logger.info(f"  Excel sheets processed: {excel_sheets}")
            
            logger.info(f"  Documents created: {len(docs)}")
            logger.info(f"  Numbered definitions: {numbered_defs}")
            logger.info(f"  Tables/sheets: {tables}")
            logger.info(f"  Complete definitions: {complete_defs}")
            
        except Exception as e:
            processing_stats['files_failed'] += 1
            logger.error(f"  Failed to process {filename}: {e}")
            logger.debug(f"  Error details: {traceback.format_exc()}")
    
    # Final comprehensive summary
    logger.info("\n" + "="*60)
    logger.info("COMPREHENSIVE DOCUMENT LOADING SUMMARY")
    logger.info("="*60)
    logger.info(f"Total files found: {processing_stats['total_files_found']}")
    logger.info(f"Files successfully processed: {processing_stats['files_processed']}")
    logger.info(f"Files failed: {processing_stats['files_failed']}")
    logger.info(f"Success rate: {(processing_stats['files_processed']/processing_stats['total_files_found']*100):.1f}%")
    logger.info("")
    logger.info("FILE TYPE BREAKDOWN:")
    logger.info(f"  Excel files: {processing_stats['excel_files']} (sheets: {processing_stats['excel_sheets']})")
    logger.info(f"  DOCX files: {processing_stats['docx_files']}")
    logger.info(f"  PDF files: {processing_stats['pdf_files']}")
    logger.info(f"  Text files: {processing_stats['txt_files']}")
    logger.info("")
    logger.info("CONTENT EXTRACTION SUMMARY:")
    logger.info(f"  Total documents created: {processing_stats['total_documents']}")
    logger.info(f"  Numbered definitions found: {processing_stats['numbered_definitions']}")
    logger.info(f"  Complete definitions: {processing_stats['complete_definitions']}")
    logger.info(f"  Tables/sheets extracted: {processing_stats['tables_extracted']}")
    logger.info(f"  Average docs per file: {processing_stats['total_documents']/max(processing_stats['files_processed'], 1):.1f}")
    logger.info("="*60)
    
    return all_documents

def create_optimized_chunks(documents: List[Document]) -> List[Document]:
    """Enhanced chunking with better handling for different content types"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=[
            "\n\n**",    # Definition headers
            "\n\n\n",    # Major section breaks
            "\n\n",      # Paragraph breaks
            "\n•",       # Bullet points
            "\n",        # Line breaks
            ". ",        # Sentence endings
            ": ",        # Definition separators
            "; ",        # Clause separators
            " ",         # Word boundaries
            ""           # Character level
        ]
    )
    
    final_chunks = []
    chunking_stats = {
        'input_documents': len(documents),
        'preserved_tables': 0,
        'preserved_definitions': 0,
        'smart_chunked_definitions': 0,
        'standard_chunked': 0,
        'total_output_chunks': 0
    }
    
    logger.info(f"Starting optimized chunking for {len(documents)} documents...")
    
    for doc_idx, doc in enumerate(documents):
        content_type = doc.metadata.get('content_type', '')
        element_type = doc.metadata.get('element_type', '')
        
        # Handle different content types appropriately
        if doc.metadata.get('is_embedded_table') or doc.metadata.get('is_table_content'):
            # Keep tables intact
            doc.metadata['chunk_strategy'] = 'preserved_table'
            final_chunks.append(doc)
            chunking_stats['preserved_tables'] += 1
            
        elif content_type == 'single_numbered_definition' and len(doc.page_content) <= MAX_DEFINITION_SIZE:
            # Keep single complete definitions intact
            doc.metadata['chunk_strategy'] = 'preserved_complete_definition'
            final_chunks.append(doc)
            chunking_stats['preserved_definitions'] += 1
            logger.info(f"Preserved complete definition: {doc.metadata.get('definition_number', 'unknown')}")
            
        elif content_type in ['single_numbered_definition', 'multiple_numbered_definitions'] or doc.metadata.get('is_numbered_definition'):
            # Use smart chunking for numbered definitions
            definition_chunks = smart_chunk_numbered_definitions(doc.page_content, text_splitter)
            
            for i, chunk in enumerate(definition_chunks):
                # Preserve and enhance metadata
                chunk.metadata.update(doc.metadata)
                chunk.metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(definition_chunks),
                    'chunk_strategy': 'smart_numbered_definition',
                    'source_doc_index': doc_idx
                })
                final_chunks.append(chunk)
            
            chunking_stats['smart_chunked_definitions'] += 1
            logger.info(f"Smart chunked definitions: {len(definition_chunks)} chunks")
                
        elif content_type == 'value_mapping_table' and len(doc.page_content) <= TABLE_CHUNK_SIZE:
            # Keep value mapping tables intact
            doc.metadata['chunk_strategy'] = 'preserved_value_mapping'
            final_chunks.append(doc)
            chunking_stats['preserved_tables'] += 1
            
        else:
            # Standard chunking for other content
            chunks = text_splitter.split_documents([doc])
            
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_strategy': 'standard_recursive',
                    'source_doc_index': doc_idx
                })
                final_chunks.append(chunk)
            
            chunking_stats['standard_chunked'] += 1
    
    chunking_stats['total_output_chunks'] = len(final_chunks)
    
    # Log chunking statistics
    logger.info("\n" + "="*60)
    logger.info("CHUNKING STRATEGY SUMMARY")
    logger.info("="*60)
    logger.info(f"Input documents: {chunking_stats['input_documents']}")
    logger.info(f"Output chunks: {chunking_stats['total_output_chunks']}")
    logger.info(f"Expansion ratio: {chunking_stats['total_output_chunks']/chunking_stats['input_documents']:.2f}x")
    logger.info("")
    logger.info("STRATEGY BREAKDOWN:")
    logger.info(f"  Preserved tables: {chunking_stats['preserved_tables']}")
    logger.info(f"  Preserved complete definitions: {chunking_stats['preserved_definitions']}")
    logger.info(f"  Smart chunked definitions: {chunking_stats['smart_chunked_definitions']}")
    logger.info(f"  Standard chunked: {chunking_stats['standard_chunked']}")
    logger.info("="*60)
    
    return final_chunks

def analyze_final_collection(documents: List[Document]) -> Dict[str, Any]:
    """Enhanced analysis of the final document collection with detailed logging"""
    logger.info("Performing comprehensive collection analysis...")
    
    analysis = {
        'total_documents': len(documents),
        'numbered_definitions': 0,
        'single_definitions': 0,
        'multiple_definitions': 0,
        'complete_definitions': 0,
        'partial_definitions': 0,
        'embedded_tables': 0,
        'excel_sheets': 0,
        'value_mapping_tables': 0,
        'field_specifications': 0,
        'content_types': {},
        'chunk_strategies': {},
        'definition_numbers': set(),
        'technical_names': set(),
        'table_types': {},
        'quality_metrics': {},
        'field_categories': {},
        'excel_processing_stats': {
            'total_excel_docs': 0,
            'sheets_processed': 0,
            'lookup_tables': 0,
            'value_mappings': 0
        },
        'hcfa_fields': {},
        'file_coverage': {},
        'chunk_strategy_details': {}
    }
    
    total_content_length = 0
    complexity_scores = []
    completeness_scores = []
    definition_lengths = []
    
    for doc in documents:
        metadata = doc.metadata
        content_length = len(doc.page_content)
        total_content_length += content_length
        
        # Track file coverage
        source_file = metadata.get('source_filename', 'unknown')
        analysis['file_coverage'][source_file] = analysis['file_coverage'].get(source_file, 0) + 1
        
        # Count different types with detailed tracking
        if metadata.get('is_numbered_definition'):
            analysis['numbered_definitions'] += 1
            definition_lengths.append(content_length)
            
            # Track definition numbers
            if metadata.get('definition_number'):
                analysis['definition_numbers'].add(metadata['definition_number'])
            
            # Track specific HCFA fields
            if metadata.get('is_hcfa_field'):
                field_type = metadata.get('hcfa_field_type', 'unknown')
                analysis['hcfa_fields'][field_type] = analysis['hcfa_fields'].get(field_type, 0) + 1
            
            # Categorize by completeness
            if metadata.get('is_complete_definition'):
                analysis['complete_definitions'] += 1
            else:
                analysis['partial_definitions'] += 1
        
        # Track content type categories
        content_type = metadata.get('content_type', 'unknown')
        analysis['content_types'][content_type] = analysis['content_types'].get(content_type, 0) + 1
        
        if content_type == 'single_numbered_definition':
            analysis['single_definitions'] += 1
        elif content_type == 'multiple_numbered_definitions':
            analysis['multiple_definitions'] += 1
        
        # Track tables and Excel data with more detail
        if metadata.get('is_embedded_table'):
            analysis['embedded_tables'] += 1
            
            table_type = metadata.get('table_type', 'unknown')
            analysis['table_types'][table_type] = analysis['table_types'].get(table_type, 0) + 1
            
            if content_type == 'value_mapping_table':
                analysis['value_mapping_tables'] += 1
            
            # Track Excel-specific stats
            if metadata.get('element_type') == 'excel_sheet':
                analysis['excel_sheets'] += 1
                analysis['excel_processing_stats']['sheets_processed'] += 1
                if metadata.get('is_lookup_table'):
                    analysis['excel_processing_stats']['lookup_tables'] += 1
                if metadata.get('is_value_mapping'):
                    analysis['excel_processing_stats']['value_mappings'] += 1
        
        # Track Excel documents
        if 'excel' in metadata.get('element_type', '').lower():
            analysis['excel_processing_stats']['total_excel_docs'] += 1
        
        # Track field specifications
        if metadata.get('is_field_spec'):
            analysis['field_specifications'] += 1
            
            # Track technical names
            if metadata.get('primary_technical_name'):
                analysis['technical_names'].add(metadata['primary_technical_name'])
            
            # Track field categories
            field_category = metadata.get('field_category', 'unknown')
            analysis['field_categories'][field_category] = analysis['field_categories'].get(field_category, 0) + 1
        
        # Track chunk strategies with more detail
        chunk_strategy = metadata.get('chunk_strategy', 'unknown')
        analysis['chunk_strategies'][chunk_strategy] = analysis['chunk_strategies'].get(chunk_strategy, 0) + 1
        
        # Detailed chunk strategy analysis
        if chunk_strategy not in analysis['chunk_strategy_details']:
            analysis['chunk_strategy_details'][chunk_strategy] = {
                'count': 0,
                'total_length': 0,
                'avg_length': 0,
                'with_definitions': 0,
                'with_tables': 0
            }
        
        strategy_detail = analysis['chunk_strategy_details'][chunk_strategy]
        strategy_detail['count'] += 1
        strategy_detail['total_length'] += content_length
        if metadata.get('is_numbered_definition'):
            strategy_detail['with_definitions'] += 1
        if metadata.get('is_embedded_table'):
            strategy_detail['with_tables'] += 1
        
        # Collect metrics
        complexity = metadata.get('complexity_score', 0)
        if isinstance(complexity, (int, float)):
            complexity_scores.append(complexity)
        
        completeness = metadata.get('definition_completeness', 0)
        if isinstance(completeness, (int, float)):
            completeness_scores.append(completeness)
    
    # Calculate average lengths for chunk strategies
    for strategy, details in analysis['chunk_strategy_details'].items():
        if details['count'] > 0:
            details['avg_length'] = details['total_length'] / details['count']
    
    # Calculate enhanced quality metrics
    analysis['quality_metrics'] = {
        'average_content_length': total_content_length / len(documents) if documents else 0,
        'average_complexity': sum(complexity_scores) / len(complexity_scores) if complexity_scores else 0,
        'average_completeness': sum(completeness_scores) / len(completeness_scores) if completeness_scores else 0,
        'definition_coverage': len(analysis['definition_numbers']),
        'complete_definition_ratio': analysis['complete_definitions'] / max(analysis['numbered_definitions'], 1),
        'table_coverage': analysis['embedded_tables'],
        'excel_coverage': analysis['excel_sheets'],
        'technical_name_coverage': len(analysis['technical_names']),
        'content_diversity': len(analysis['content_types']),
        'file_coverage': len(analysis['file_coverage']),
        'average_definition_length': sum(definition_lengths) / len(definition_lengths) if definition_lengths else 0,
        'hcfa_field_coverage': len(analysis['hcfa_fields'])
    }
    
    # Convert sets to lists for JSON serialization
    analysis['definition_numbers'] = sorted(list(analysis['definition_numbers']))
    analysis['technical_names'] = sorted(list(analysis['technical_names']))
    
    # Log comprehensive analysis results
    logger.info("\n" + "="*80)
    logger.info("COMPREHENSIVE COLLECTION ANALYSIS RESULTS")
    logger.info("="*80)
    logger.info(f"Total chunks: {analysis['total_documents']}")
    logger.info("")
    logger.info("NUMBERED DEFINITIONS:")
    logger.info(f"  Total: {analysis['numbered_definitions']}")
    logger.info(f"  Single: {analysis['single_definitions']}")
    logger.info(f"  Multiple: {analysis['multiple_definitions']}")
    logger.info(f"  Complete: {analysis['complete_definitions']} ({analysis['quality_metrics']['complete_definition_ratio']:.1%})")
    logger.info(f"  Unique numbers: {analysis['quality_metrics']['definition_coverage']}")
    logger.info("")
    logger.info("TABLES AND DATA:")
    logger.info(f"  Embedded tables: {analysis['embedded_tables']}")
    logger.info(f"  Excel sheets: {analysis['excel_sheets']}")
    logger.info(f"  Value mapping tables: {analysis['value_mapping_tables']}")
    logger.info(f"  Field specifications: {analysis['field_specifications']}")
    logger.info("")
    logger.info("QUALITY METRICS:")
    logger.info(f"  Average content length: {analysis['quality_metrics']['average_content_length']:.0f} chars")
    logger.info(f"  Average definition length: {analysis['quality_metrics']['average_definition_length']:.0f} chars")
    logger.info(f"  Average completeness: {analysis['quality_metrics']['average_completeness']:.1f}/5")
    logger.info(f"  Content diversity: {analysis['quality_metrics']['content_diversity']} types")
    logger.info(f"  Technical names: {analysis['quality_metrics']['technical_name_coverage']}")
    logger.info("")
    logger.info("FILE COVERAGE:")
    for filename, count in sorted(analysis['file_coverage'].items()):
        logger.info(f"  {filename}: {count} chunks")
    logger.info("")
    logger.info("CONTENT TYPES:")
    for content_type, count in sorted(analysis['content_types'].items()):
        percentage = (count / analysis['total_documents']) * 100
        logger.info(f"  {content_type}: {count} ({percentage:.1f}%)")
    logger.info("")
    logger.info("CHUNK STRATEGIES:")
    for strategy, count in sorted(analysis['chunk_strategies'].items()):
        percentage = (count / analysis['total_documents']) * 100
        details = analysis['chunk_strategy_details'][strategy]
        logger.info(f"  {strategy}: {count} ({percentage:.1f}%) - avg length: {details['avg_length']:.0f}")
    
    if analysis['hcfa_fields']:
        logger.info("")
        logger.info("HCFA FIELDS:")
        for field_type, count in sorted(analysis['hcfa_fields'].items()):
            logger.info(f"  {field_type}: {count}")
    
    logger.info("="*80)
    
    return analysis

def main():
    """Enhanced main function with comprehensive logging and error handling"""
    logger.info("="*80)
    logger.info("STARTING COMPREHENSIVE CHROMADB CREATION")
    logger.info("Enhanced for numbered definitions, embedded tables, and XLS processing")
    logger.info("="*80)
    
    start_time = time.time()
    
    # Clean existing database
    if os.path.exists(CHROMA_DB_DIR):
        logger.info(f"Removing existing ChromaDB at {CHROMA_DB_DIR}")
        shutil.rmtree(CHROMA_DB_DIR)
    
    os.makedirs(CHROMA_DB_DIR, exist_ok=True)
    logger.info(f"Created ChromaDB directory: {CHROMA_DB_DIR}")
    
    # Load and process documents
    logger.info("="*60)
    logger.info("PHASE 1: DOCUMENT LOADING AND EXTRACTION")
    logger.info("="*60)
    documents = load_all_documents(DOCUMENTS_DIR)
    
    if not documents:
        logger.error("No documents loaded. Check source directory and file formats.")
        logger.error(f"Source directory: {DOCUMENTS_DIR}")
        logger.error(f"Supported formats: .docx, .doc, .xlsx, .xls, .pdf, .txt")
        return
    
    # Create optimized chunks
    logger.info("="*60)
    logger.info("PHASE 2: INTELLIGENT CHUNKING")
    logger.info("="*60)
    chunked_docs = create_optimized_chunks(documents)
    
    # Analyze the collection
    logger.info("="*60)
    logger.info("PHASE 3: COLLECTION ANALYSIS")
    logger.info("="*60)
    analysis = analyze_final_collection(chunked_docs)
    
    # Prepare data for ChromaDB
    logger.info("="*60)
    logger.info("PHASE 4: EMBEDDING GENERATION")
    logger.info("="*60)
    
    ids = [f"doc_{i}" for i in range(len(chunked_docs))]
    contents = [doc.page_content for doc in chunked_docs]
    metadatas = [doc.metadata for doc in chunked_docs]
    
    # Generate embeddings with progress tracking
    logger.info(f"Generating embeddings for {len(contents)} documents...")
    try:
        # Process in smaller batches for large datasets
        batch_size = 50
        embeddings_list = []
        
        for i in range(0, len(contents), batch_size):
            batch_contents = contents[i:i+batch_size]
            batch_embeddings = embeddings.embed_documents(batch_contents)
            embeddings_list.extend(batch_embeddings)
            
            progress = (i + len(batch_contents)) / len(contents) * 100
            logger.info(f"  Embedding progress: {progress:.1f}% ({i + len(batch_contents)}/{len(contents)})")
        
        logger.info(f"✓ Generated {len(embeddings_list)} embeddings successfully")
        
    except Exception as e:
        logger.error(f"Error generating embeddings: {e}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return
    
    # Initialize ChromaDB and store documents
    logger.info("="*60)
    logger.info("PHASE 5: CHROMADB STORAGE")
    logger.info("="*60)
    
    try:
        client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
        collection = client.get_or_create_collection(name=COLLECTION_NAME)
        logger.info(f"Created/connected to collection: {COLLECTION_NAME}")
        
        # Store in batches to handle large datasets
        batch_size = 500
        for i in range(0, len(ids), batch_size):
            end_idx = min(i + batch_size, len(ids))
            batch_ids = ids[i:end_idx]
            batch_embeddings = embeddings_list[i:end_idx]
            batch_contents = contents[i:end_idx]
            batch_metadatas = metadatas[i:end_idx]
            
            collection.add(
                ids=batch_ids,
                embeddings=batch_embeddings,
                documents=batch_contents,
                metadatas=batch_metadatas
            )
            
            progress = end_idx / len(ids) * 100
            logger.info(f"  Storage progress: {progress:.1f}% ({end_idx}/{len(ids)})")
        
        # Verify final count
        final_count = collection.count()
        logger.info(f"✓ Final collection size: {final_count} documents")
        
        # Comprehensive testing
        logger.info("="*60)
        logger.info("PHASE 6: COMPREHENSIVE TESTING")
        logger.info("="*60)
        
        # Enhanced test queries with expected results
        test_queries = [
            ("141. HCFA Admit Type Code", "Should find exact definition with all fields"),
            ("HCFA Admit Type Code", "Should find field definition by name"),
            ("Diagnosis Code 6", "Should find diagnosis field definition"),
            ("Technical Name: HCFA_ADMIT_TYPE_CD", "Should find by technical name"),
            ("158. Diagnosis Code 6", "Should find specific diagnosis field"),
            ("Field 141", "Should find field by number reference")
        ]
        
        for query, expected in test_queries:
            try:
                logger.info(f"\nTesting query: '{query}'")
                logger.info(f"Expected: {expected}")
                
                test_embedding = embeddings.embed_query(query)
                test_results = collection.query(
                    query_embeddings=[test_embedding],
                    n_results=5,
                    include=["documents", "metadatas", "distances"]
                )
                
                if test_results and test_results.get('documents'):
                    results_count = len(test_results['documents'][0])
                    logger.info(f"  ✓ Found {results_count} results")
                    
                    for i in range(min(3, results_count)):  # Show top 3 results
                        doc = test_results['documents'][0][i]
                        metadata = test_results.get('metadatas', [[]])[0][i] if test_results.get('metadatas') else {}
                        distance = test_results.get('distances', [[]])[0][i] if test_results.get('distances') else 1.0
                        
                        logger.info(f"    Result {i+1}:")
                        logger.info(f"      Distance: {distance:.4f}")
                        logger.info(f"      Content preview: {doc[:100]}...")
                        
                        # Check for specific matches
                        if metadata.get('is_numbered_definition'):
                            def_num = metadata.get('definition_number', 'N/A')
                            def_title = metadata.get('definition_title', 'N/A')
                            logger.info(f"      Found numbered definition: {def_num}. {def_title}")
                        
                        if metadata.get('is_embedded_table'):
                            table_type = metadata.get('table_type', 'unknown')
                            logger.info(f"      Found table: {table_type}")
                else:
                    logger.warning(f"  ✗ No results found for query '{query}'")
                
            except Exception as e:
                logger.error(f"  ✗ Test query '{query}' failed: {e}")
        
        # Additional specific tests for numbered definitions
        logger.info("\n--- Testing Numbered Definition Retrieval ---")
        numbered_def_tests = [
            "141",  # Just the number
            "158",  # Another specific number
            "100",  # Common field number
        ]
        
        for number in numbered_def_tests:
            try:
                # Search for documents with specific definition numbers
                where_clause = {"definition_number": number}
                
                # Query with metadata filter
                test_results = collection.query(
                    query_texts=[f"Field {number}"],
                    n_results=3,
                    where=where_clause,
                    include=["documents", "metadatas", "distances"]
                )
                
                if test_results and test_results.get('documents') and test_results['documents'][0]:
                    logger.info(f"  ✓ Found definition {number} via metadata filter")
                    doc = test_results['documents'][0][0]
                    logger.info(f"    Preview: {doc[:150]}...")
                else:
                    logger.warning(f"  ✗ No definition found for number {number}")
                    
            except Exception as e:
                logger.error(f"  ✗ Metadata test for {number} failed: {e}")
        
        # Test HCFA-specific fields
        logger.info("\n--- Testing HCFA Field Retrieval ---")
        hcfa_tests = [
            ("is_hcfa_field", True),
            ("field_category", "hcfa"),
        ]
        
        for key, value in hcfa_tests:
            try:
                test_results = collection.query(
                    query_texts=["HCFA"],
                    n_results=3,
                    where={key: value},
                    include=["documents", "metadatas", "distances"]
                )
                
                if test_results and test_results.get('documents') and test_results['documents'][0]:
                    count = len(test_results['documents'][0])
                    logger.info(f"  ✓ Found {count} HCFA fields with {key}={value}")
                else:
                    logger.warning(f"  ✗ No HCFA fields found with {key}={value}")
                    
            except Exception as e:
                logger.error(f"  ✗ HCFA test failed: {e}")
        
        # Save analysis report
        logger.info("="*60)
        logger.info("PHASE 7: SAVING ANALYSIS REPORT")
        logger.info("="*60)
        
        analysis_file = os.path.join(CHROMA_DB_DIR, "collection_analysis.json")
        try:
            with open(analysis_file, 'w') as f:
                json.dump(analysis, f, indent=2, default=str)
            logger.info(f"✓ Saved collection analysis to {analysis_file}")
        except Exception as e:
            logger.error(f"✗ Failed to save analysis report: {e}")
        
        # Create a summary report
        summary_file = os.path.join(CHROMA_DB_DIR, "creation_summary.txt")
        try:
            with open(summary_file, 'w') as f:
                f.write("ChromaDB Creation Summary\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Creation Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Processing Time: {time.time() - start_time:.2f} seconds\n\n")
                
                f.write("Collection Statistics:\n")
                f.write(f"  Total Documents: {final_count}\n")
                f.write(f"  Numbered Definitions: {analysis['numbered_definitions']}\n")
                f.write(f"  Complete Definitions: {analysis['complete_definitions']}\n")
                f.write(f"  Excel Sheets: {analysis['excel_sheets']}\n")
                f.write(f"  Embedded Tables: {analysis['embedded_tables']}\n")
                f.write(f"  Technical Names: {len(analysis['technical_names'])}\n")
                f.write(f"  Definition Numbers: {len(analysis['definition_numbers'])}\n\n")
                
                f.write("Key Features:\n")
                f.write("  ✓ Enhanced numbered definition detection\n")
                f.write("  ✓ Comprehensive XLS/Excel processing\n")
                f.write("  ✓ Smart chunking for complete definitions\n")
                f.write("  ✓ Advanced metadata enrichment\n")
                f.write("  ✓ HCFA field specialized handling\n")
                f.write("  ✓ Table and appendix extraction\n")
                
            logger.info(f"✓ Saved creation summary to {summary_file}")
        except Exception as e:
            logger.error(f"✗ Failed to save summary report: {e}")
        
        # Final success message
        end_time = time.time()
        processing_time = end_time - start_time
        
        logger.info("\n" + "="*80)
        logger.info("✅ CHROMADB CREATION COMPLETED SUCCESSFULLY! ✅")
        logger.info("="*80)
        logger.info("Key Achievements:")
        logger.info(f"  🔢 {final_count} total documents indexed")
        logger.info(f"  📋 {analysis['numbered_definitions']} numbered definitions processed")
        logger.info(f"  📊 {analysis['excel_sheets']} Excel sheets processed")
        logger.info(f"  📑 {analysis['embedded_tables']} tables/appendices extracted")
        logger.info(f"  ✅ {analysis['complete_definitions']} complete definitions preserved")
        logger.info(f"  🏥 {len(analysis['hcfa_fields'])} HCFA field types covered")
        logger.info(f"  🔍 {len(analysis['definition_numbers'])} unique definition numbers")
        logger.info(f"  ⚡ {processing_time:.2f} seconds total processing time")
        logger.info("")
        logger.info("System Optimizations:")
        logger.info("  • Enhanced pattern matching for numbered definitions")
        logger.info("  • Smart chunking preserves complete field definitions")
        logger.info("  • Comprehensive XLS/table processing with value mappings")
        logger.info("  • Rich metadata for precise retrieval")
        logger.info("  • HCFA field specialized classification")
        logger.info("  • Robust error handling and fallback mechanisms")
        logger.info("")
        logger.info("The system is now ready for queries like:")
        logger.info("  • '141. HCFA Admit Type Code'")
        logger.info("  • 'HCFA Admit Type Code'")
        logger.info("  • 'Technical Name: HCFA_ADMIT_TYPE_CD'")
        logger.info("  • 'Field 141'")
        logger.info("="*80)
        
    except Exception as e:
        logger.error(f"Error creating ChromaDB: {e}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        
        # Try to provide helpful error information
        logger.error("\nTroubleshooting suggestions:")
        logger.error("1. Ensure Google Cloud authentication is properly configured")
        logger.error("2. Verify Vertex AI API is enabled and accessible")
        logger.error("3. Check that the documents directory contains valid files")
        logger.error("4. Ensure sufficient disk space for ChromaDB storage")
        logger.error("5. Review the full error log above for specific issues")

if __name__ == "__main__":
    # Check prerequisites
    if not os.path.exists(DOCUMENTS_DIR):
        logger.error(f"Documents directory '{DOCUMENTS_DIR}' not found")
        logger.info("Creating documents directory...")
        os.makedirs(DOCUMENTS_DIR, exist_ok=True)
        logger.warning("Please add document files to this directory and run again")
        logger.info(f"Supported formats: .docx, .doc, .xlsx, .xls, .pdf, .txt")
    elif not os.listdir(DOCUMENTS_DIR):
        logger.warning(f"Documents directory '{DOCUMENTS_DIR}' is empty")
        logger.info("Please add document files to this directory")
        logger.info(f"Supported formats: .docx, .doc, .xlsx, .xls, .pdf, .txt")
    else:
        # Run the main process
        main()